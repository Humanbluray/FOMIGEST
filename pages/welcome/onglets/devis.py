from utils import *
import flet as ft
import backend as be
from utils.useful_functions import ajout_separateur, ecrire_en_lettres, ecrire_date
import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import io
from dotenv import load_dotenv
from supabase import create_client

load_dotenv()
url = os.environ.get("SUPABASE_URL")
key = os.environ.get("SUPABASE_KEY")
supabase = create_client(url, key)


class OneArticle(ft.Container):
    def __init__(self, cp: object, ref: str, des: str, prix: int):
        super().__init__(
            padding=ft.padding.only(10, 3, 10, 3),
        )
        self.cp = cp
        self.ref = ref
        self.des = des
        self.prix = ft.TextField(**numbers_field_style, value="", width=100, label="Prix")
        self.qte = ft.TextField(**numbers_field_style, value="", width=80, label="Qté")
        self.content = ft.Row(
            controls=[
                ft.Row(
                    controls=[
                        ft.TextField(**readonly_field_style, value=ref, width=170, label="Référence", tooltip=f"Prix: {ajout_separateur(prix)}"),
                        ft.TextField(**readonly_field_style, value=des, width=300, label="Nom pièce"),
                        self.qte, self.prix
                    ]
                ),
                CtButton("add", ft.colors.BLACK45,"Ajouter", None, self.add_article_to_list)
            ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
        )

    def add_article_to_list(self, e):
        count = 0
        for widget in (self.qte, self.prix):
            if widget.value is None or widget.value == "":
                count += 1

        if count == 0:
            self.cp.new_table.controls.append(
                NewOneArticle(self.cp, self.ref, self.des, int(self.qte.value), int(self.prix.value))
            )
            self.cp.new_table.update()
            self.cp.cp.box.title.value = "Validé"
            self.cp.cp.box.content.value = "Article ajouté"
            self.cp.cp.box.open = True
            self.cp.cp.box.update()

            somme = 0
            nb_ref = 0
            for widget in self.cp.new_table.controls[:]:
                somme += (int(widget.art_prix.value) * int(widget.art_qte.value))
                nb_ref += 1

            self.cp.new_montant.value = somme
            total = int(somme - (somme * int(self.cp.new_remise.value) / 100))
            self.cp.new_lettres.value = ecrire_en_lettres(total)
            self.cp.new_nb_ref.value = f"{nb_ref} article(s)"

            for widget in (self.cp.new_nb_ref, self.cp.new_montant, self.cp.new_lettres):
                widget.update()

            for widget in (self.prix, self.qte):
                widget.value = None
                widget.update()

        else:
            self.cp.cp.box.title.value = "Erreur"
            self.cp.cp.box.content.value = "Qté et prix sont obligatoires"
            self.cp.cp.box.open = True
            self.cp.cp.box.update()


class AddOneArticle(ft.Container):
    def __init__(self, cp: object, ref: str, des: str, prix: int):
        super().__init__(
            padding=ft.padding.only(10, 3, 10, 3),
        )
        self.cp = cp
        self.ref = ref
        self.des = des
        self.prix = ft.TextField(**numbers_field_style, value="", width=100, label="Prix")
        self.qte = ft.TextField(**numbers_field_style, value="", width=80, label="Qté")
        self.content = ft.Row(
            controls=[
                ft.Row(
                    controls=[
                        ft.TextField(**readonly_field_style, value=ref, width=170, label="Référence", tooltip=f"Prix: {ajout_separateur(prix)}"),
                        ft.TextField(**readonly_field_style, value=des, width=300, label="Nom pièce"),
                        self.qte, self.prix
                    ]
                ),
                CtButton("add", ft.colors.BLACK45, "Ajouter", None, self.add_article_to_list)
            ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
        )

    def add_article_to_list(self, e):
        count = 0
        for widget in (self.qte, self.prix):
            if widget.value is None or widget.value == "":
                count += 1

        if count == 0:
            self.cp.edit_table.controls.append(
                EditOneArticle(self.cp, self.ref, self.des, int(self.qte.value), int(self.prix.value))
            )
            self.cp.edit_table.update()
            self.cp.cp.box.title.value = "Validé"
            self.cp.cp.box.content.value = "Article ajouté"
            self.cp.cp.box.open=True
            self.cp.cp.box.update()

            somme = 0
            nb_ref = 0
            for widget in self.cp.edit_table.controls[:]:
                somme += (int(widget.art_prix.value) * int(widget.art_qte.value))
                nb_ref += 1

            self.cp.edit_montant.value = somme
            total = int(somme - (somme * int(self.cp.edit_remise.value) / 100))
            self.cp.edit_lettres.value = ecrire_en_lettres(total)
            self.cp.edit_nb_ref.value = f"{nb_ref} article(s)"

            for widget in (self.cp.edit_nb_ref, self.cp.edit_montant, self.cp.edit_lettres):
                widget.update()

            for widget in (self.prix, self.qte):
                widget.value = None
                widget.update()

        else:
            self.cp.cp.box.title.value = "Erreur"
            self.cp.cp.box.content.value = "Qté et prix sont obligatoires"
            self.cp.cp.box.open = True
            self.cp.cp.box.update()


class EditOneArticle(ft.Container):
    def __init__(self, cp: object, ref: str, des: str, qte: int, prix: int):
        super().__init__(
            padding=ft.padding.only(10, 3, 10, 3),
        )
        self.cp = cp
        self.ref = ref
        self.des = des
        self.prix = prix
        self.qte = qte
        self.status = False
        self.art_ref = ft.TextField(**readonly_field_style, value=ref, width=140, label="Référence",)
        self.art_des = ft.TextField(**readonly_field_style, value=des, width=300, label="Nom pièce", tooltip=f"{des}")
        self.art_qte = ft.TextField(**numbers_field_style, value=f"{qte}", width=80, label="Qté", disabled=True)
        self.art_prix = ft.TextField(**numbers_field_style, value=f"{prix}", width=100, label="Prix", disabled=True)
        self.bt_modif = CtButton(ft.icons.CHECK, ft.colors.BLACK45, "Modifier", None, self.edit_article)
        self.bt_modif.visible = False
        self.bt_allow_modif = CtButton("edit_outlined", ft.colors.BLACK45, "Modifier", None, self.allow_edit)
        self.bt_allow_modif.visible = True
        self.content = ft.Row(
            controls=[
                ft.Row(
                    controls=[self.art_ref, self.art_des, self.art_qte, self.art_prix]
                ),
               ft.Row(
                   controls=[
                       self.bt_modif, self.bt_allow_modif,
                       CtButton("delete_outlined", ft.colors.BLACK45, "Modifier", None, self.delete_article)
                   ], spacing=0
               )
            ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
        )

    def allow_edit(self, e):
        # if not self.status == "edit_outlined":
        self.bt_modif.visible = True
        self.bt_modif.update()
        e.control.visible = False
        e.control.update()
        self.status = True
        for widget in (self.art_qte, self.art_prix):
            widget.disabled = False
            widget.update()

    def edit_article(self, e):
        e.control.visible = False
        e.control.update()
        self.bt_allow_modif.visible = True
        self.bt_allow_modif.update()
        self.status = False
        for widget in (self.art_qte, self.art_prix):
            widget.disabled = True
            widget.update()

        somme = 0
        nb_ref = 0

        if len(self.cp.edit_table.controls[:]) == 0:
            self.cp.edit_montant.value = "0"
            self.cp.edit_lettres.value = ecrire_en_lettres(0)
            self.cp.edit_nb_ref.value = f"Aucune ligne(s)"

        else:
            for widget in self.cp.edit_table.controls[:]:
                somme += (int(widget.art_prix.value) * int(widget.art_qte.value))
                nb_ref += 1

            self.cp.edit_montant.value = f"{somme}"
            total = int(somme - (somme * int(self.cp.edit_remise.value) / 100))
            self.cp.edit_lettres.value = ecrire_en_lettres(total)
            self.cp.edit_nb_ref.value = f"{nb_ref} ligne(s)"

        for widget in (self.cp.edit_nb_ref, self.cp.edit_montant, self.cp.edit_lettres):
            widget.update()

        self.cp.edit_table.update()

    def delete_article(self, e):
        self.cp.edit_table.controls.remove(self)
        somme = 0
        nb_ref = 0

        if len(self.cp.edit_table.controls[:]) == 0:
            self.cp.edit_montant.value = "0"
            self.cp.edit_lettres.value = ecrire_en_lettres(0)
            self.cp.edit_nb_ref.value = f"Aucune ligne(s)"

        else:
            for widget in self.cp.edit_table.controls[:]:
                somme += (int(widget.art_prix.value) * int(widget.art_qte.value))
                nb_ref += 1

            self.cp.edit_montant.value = f"{somme}"
            total = int(somme - (somme * int(self.cp.edit_remise.value) / 100))
            self.cp.edit_lettres.value = ecrire_en_lettres(total)
            self.cp.edit_nb_ref.value = f"{nb_ref} ligne(s)"

        for widget in (self.cp.edit_nb_ref, self.cp.edit_montant, self.cp.edit_lettres):
            widget.update()

        self.cp.edit_table.update()


class NewOneArticle(ft.Container):
    def __init__(self, cp: object, ref: str, des: str, qte: int, prix: int):
        super().__init__(
            padding=ft.padding.only(10, 3, 10, 3),
        )
        self.cp = cp
        self.ref = ref
        self.des = des
        self.prix = prix
        self.qte = qte
        self.status = False
        self.art_ref = ft.TextField(**readonly_field_style, value=ref, width=140, label="Référence",)
        self.art_des = ft.TextField(**readonly_field_style, value=des, width=300, label="Nom pièce", tooltip=f"{des}")
        self.art_qte = ft.TextField(**numbers_field_style, value=f"{qte}", width=80, label="Qté", disabled=True)
        self.art_prix = ft.TextField(**numbers_field_style, value=f"{prix}", width=100, label="Prix", disabled=True)
        self.bt_modif = CtButton(ft.icons.CHECK, ft.colors.BLACK45, "Modifier", None, self.edit_article)
        self.bt_modif.visible = False
        self.bt_allow_modif = CtButton("edit_outlined", "", "Modifier", None, self.allow_edit)
        self.bt_allow_modif.visible = True
        self.content = ft.Row(
            controls=[
                ft.Row(
                    controls=[self.art_ref, self.art_des, self.art_qte, self.art_prix]
                ),
               ft.Row(
                   controls=[
                       self.bt_modif, self.bt_allow_modif,
                       CtButton("delete_outlined", ft.colors.RED,"Modifier", None, self.delete_article)
                   ], spacing=0
               )
            ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
        )

    def allow_edit(self, e):
        # if not self.status == "edit_outlined":
        self.bt_modif.visible = True
        self.bt_modif.update()
        e.control.visible = False
        e.control.update()
        self.status = True
        for widget in (self.art_qte, self.art_prix):
            widget.disabled = False
            widget.update()

    def edit_article(self, e):
        e.control.visible = False
        e.control.update()
        self.bt_allow_modif.visible = True
        self.bt_allow_modif.update()
        self.status = False
        for widget in (self.art_qte, self.art_prix):
            widget.disabled = True
            widget.update()

        somme = 0
        nb_ref = 0

        if len(self.cp.new_table.controls[:]) == 0:
            self.cp.new_montant.value = "0"
            self.cp.new_lettres.value = ecrire_en_lettres(0)
            self.cp.new_nb_ref.value = f"Aucune ligne(s)"

        else:
            for widget in self.cp.new_table.controls[:]:
                somme += (int(widget.art_prix.value) * int(widget.art_qte.value))
                nb_ref += 1

            self.cp.new_montant.value = f"{somme}"
            total = int(somme - (somme * int(self.cp.new_remise.value) / 100))
            self.cp.new_lettres.value = ecrire_en_lettres(total)
            self.cp.new_nb_ref.value = f"{nb_ref} ligne(s)"

        for widget in (self.cp.new_nb_ref, self.cp.new_montant, self.cp.new_lettres):
            widget.update()

        self.cp.new_table.update()

    def delete_article(self, e):
        self.cp.new_table.controls.remove(self)
        somme = 0
        nb_ref = 0

        if len(self.cp.new_table.controls[:]) == 0:
            self.cp.new_montant.value = "0"
            self.cp.new_lettres.value = ecrire_en_lettres(0)
            self.cp.new_nb_ref.value = f"Aucune ligne(s)"

        else:
            for widget in self.cp.new_table.controls[:]:
                somme += (int(widget.art_prix.value) * int(widget.art_qte.value))
                nb_ref += 1

            self.cp.new_montant.value = f"{somme}"
            total = int(somme - (somme * int(self.cp.new_remise.value) / 100))
            self.cp.new_lettres.value = ecrire_en_lettres(total)
            self.cp.new_nb_ref.value = f"{nb_ref} ligne(s)"

        for widget in (self.cp.new_nb_ref, self.cp.new_montant, self.cp.new_lettres):
            widget.update()

        self.cp.new_table.update()


class Devis(ft.Container):
    def __init__(self, cp: object):
        super().__init__(expand=True)
        self.cp = cp
        self.nb_devis = ft.Text("", size=18, font_family="Poppins Medium")
        self.nb_devis_expires = ft.Text("", size=18, font_family="Poppins Medium")
        self.nb_devis_encours = ft.Text("", size=18, font_family="Poppins Medium")
        self.search = ft.TextField(**search_field_style, width=300, prefix_icon="search", on_change=self.filter_datas)
        self.results = ft.Text("", size=12, font_family="Poppins Medium")
        self.table = ft.DataTable(
            **datatable_style,
            columns=[
                ft.DataColumn(ft.Text("")),
                ft.DataColumn(ft.Text("Date".upper())),
                ft.DataColumn(ft.Text("numero".upper())),
                ft.DataColumn(ft.Text("client".upper())),
                ft.DataColumn(ft.Text("Montant".upper())),
                ft.DataColumn(ft.Text("Actions".upper())),
            ]
        )
        self.main_window = ft.Container(
            expand=True,
            padding=ft.padding.only(20, 15, 20, 15), border_radius=10, bgcolor="white",
            content=ft.Column(
                controls=[
                    ft.Container(
                        bgcolor="#f0f0f6", padding=10, border_radius=16,
                        content=ft.Row(
                            controls=[
                                ft.Icon(ft.icons.MONETIZATION_ON, color="black"),
                                ft.Text("Devis".upper(), size=24, font_family="Poppins Bold"),
                            ]
                        )
                    ),
                    ft.Divider(height=1, color=ft.colors.TRANSPARENT),
                    ft.Row(
                        controls=[
                            ft.Row(
                                controls=[
                                    self.search, self.results
                                ]
                            ),
                            AnyButton(FIRST_COLOR, "add", "Créer devis", "white", 175, self.open_new_window)
                        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
                    ),
                    ft.Divider(height=1, color=ft.colors.TRANSPARENT),
                    ft.ListView(expand=True, controls=[self.table])
                ]
            )
        )

        # New devis window ...
        self.new_devis_num = ft.Text("", size=12, font_family="Poppins Medium")
        self.new_objet = ft.TextField(
            **field_style, width=300, label="Objet du devis",
        )
        self.new_client = ft.Dropdown(**drop_style, label="Client", width=400, prefix_icon=ft.icons.PERSON_SEARCH_OUTLINED, on_change=self.on_change_new_client)
        self.new_nb_ref = ft.Text("", size=12, font_family="Poppins Medium", italic=True, color="grey")
        self.new_table = ft.ListView(expand=True, divider_thickness=1, spacing=10)
        self.new_num = ft.Text("", size=12, font_family="Poppins Medium", color="white", )
        self.new_montant = ft.TextField(**readonly_field_style, width=140, prefix_icon=ft.icons.PRICE_CHANGE_OUTLINED,
                                         label="Montant")
        self.new_lettres = ft.TextField(**readonly_field_style, width=700, prefix_icon=ft.icons.LABEL_OUTLINED,
                                         label="Montant en lettres")
        self.new_objet = ft.TextField(**field_style, width=700, label="Objet du devis",
                                       prefix_icon=ft.icons.MAIL_OUTLINED)
        self.new_notabene = ft.TextField(**field_style, width=700, label="NB", prefix_icon=ft.icons.NOTES_OUTLINED)
        self.new_paiement = ft.TextField(**numbers_field_style, width=110, label="Paiement",
                                          prefix_icon=ft.icons.WATCH_LATER_OUTLINED)
        self.new_validite = ft.TextField(**numbers_field_style, width=120, label="validité",
                                          prefix_icon=ft.icons.EVENT_AVAILABLE)
        self.new_point_liv = ft.TextField(**field_style, width=200, label="pt. livraison",
                                           prefix_icon=ft.icons.LOCATION_ON_OUTLINED)
        self.new_delai = ft.TextField(**field_style, width=150, label="Délai livraison",
                                       prefix_icon=ft.icons.TIMELAPSE_OUTLINED)
        self.bt_dim_new_rem = MiniCtButton(ft.icons.KEYBOARD_ARROW_LEFT_OUTLINED, "Réduire", None, self.put_new_remise_down)
        self.bt_add_new_rem = MiniCtButton(ft.icons.KEYBOARD_ARROW_RIGHT_OUTLINED, "AUgmenter", None, self.put_new_remise_up)
        self.new_remise = ft.TextField(**readonly_field_style, width=60, label="remise", value="0")
        self.bt_create_dev = AnyButton(FIRST_COLOR, "check", "Créer devis", "white", 180, self.create_new_devis)
        self.new_window = ft.Card(
            elevation=20, surface_tint_color="#f0f0f6", width=900, height=660,
            clip_behavior=ft.ClipBehavior.ANTI_ALIAS, shadow_color="black",
            scale=ft.transform.Scale(0), expand=True,
            animate_scale=ft.Animation(300, ft.AnimationCurve.DECELERATE),
            content=ft.Container(
                padding=10, bgcolor="#f0f0f6", expand=True,
                content=ft.Column(
                    expand=True, scroll=ft.ScrollMode.AUTO,
                    controls=[
                        ft.Container(
                            bgcolor="white", padding=ft.padding.only(10, 5, 10, 5), border_radius=16,
                            content=ft.Row(
                                controls=[
                                    ft.Row(
                                        controls=[
                                            ft.Icon(ft.icons.NOTE_ADD_OUTLINED, "black"),
                                            ft.Text("Créer devis".upper(), size=14,
                                                    font_family="Poppins Bold")
                                        ]
                                    ),
                                    ft.IconButton("close", FIRST_COLOR, scale=0.6, bgcolor="#f2f2f2", on_click=self.close_new_window)
                                ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
                            )
                        ),
                        ft.Container(
                            bgcolor="white", padding=20, border_radius=16, expand=True,
                            content=ft.Column(
                                controls=[
                                    ft.Row(
                                        controls=[
                                            ft.Row(
                                                controls=[
                                                    ft.Row(
                                                        controls=[
                                                            self.new_client,
                                                            ft.Text("Devis N°".upper(), size=12, font_family="Poppins Medium"),
                                                            ft.Container(
                                                                padding=ft.padding.only(10, 3, 10, 3), border_radius=10,
                                                                bgcolor=ft.colors.DEEP_ORANGE, content=ft.Row(
                                                                    controls=[self.new_num],
                                                                    alignment=ft.MainAxisAlignment.CENTER
                                                                )
                                                            )
                                                        ]
                                                    ),
                                                ]
                                            ),
                                            self.new_nb_ref
                                        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
                                    ),
                                    ft.Stack(
                                        controls=[
                                            ft.Container(
                                                height=220, padding=ft.padding.only(10, 3, 10, 3),
                                                border_radius=16,
                                                expand=True,
                                                border=ft.border.all(1, "grey"), content=self.new_table
                                            ),
                                            ft.FloatingActionButton(
                                                bgcolor="black", on_click=self.open_new_add_ref_window,
                                                content=ft.Row([ft.Icon(ft.icons.ADD, color="white")],
                                                               alignment=ft.MainAxisAlignment.CENTER),
                                                bottom=10, scale=0.8, right=10, tooltip="Ajouter référence",
                                                opacity=0.80
                                            )
                                        ], alignment=ft.alignment.bottom_right
                                    ),
                                    ft.Row([self.new_objet]),
                                    ft.Row([self.new_paiement, self.new_delai, self.new_point_liv,
                                            self.new_validite,
                                            ft.Row([self.bt_dim_new_rem,  self.new_remise, self.bt_add_new_rem], spacing=0)
                                            ]),
                                    self.new_notabene,
                                    ft.Row([self.new_montant, self.new_lettres]),
                                    self.bt_create_dev
                                ]
                            )
                        )
                    ]
                )
            )
        )
        # Ajouter new reference window ...
        self.new_add_select_article = ft.TextField(**search_field_style, width=300, on_change=self.filter_selection_new,
                                               prefix_icon="search")
        self.new_table_add_ref = ft.ListView(expand=True, divider_thickness=1, spacing=10)
        self.new_add_ref_window = ft.Card(
            elevation=20, surface_tint_color="#f0f0f6", width=800, height=550,
            clip_behavior=ft.ClipBehavior.ANTI_ALIAS, shadow_color="black",
            scale=ft.transform.Scale(0), expand=True,
            animate_scale=ft.Animation(300, ft.AnimationCurve.DECELERATE),
            content=ft.Container(
                padding=10, bgcolor="#f0f0f6", expand=True,
                content=ft.Container(
                    expand=True, padding=10, border_radius=16, bgcolor="white",
                    content=ft.Column(
                        expand=True,
                        controls=[
                            ft.Column(
                                controls=[
                                    ft.Text("Ajouter article".upper(), size=14, font_family="Poppins Bold"),
                                    ft.Divider(height=1, thickness=1),
                                ], spacing=0
                            ),
                            self.new_add_select_article,
                            ft.Container(
                                border_radius=12, border=ft.border.all(1, "grey"), expand=True,
                                padding=ft.padding.only(10, 5, 10, 5),
                                content=self.new_table_add_ref
                            ),
                            ft.Row(
                                [AnyButton(FIRST_COLOR, "close", "Fermer", "white", 170, self.close_new_add_ref_window)],
                                alignment=ft.MainAxisAlignment.END
                            )
                        ]
                    )
                )
            )
        )

        # Edit window ...
        self.edit_bt_facture = ft.Container(padding=ft.padding.only(10, 3, 10, 3), border_radius=10)
        self.edit_nb_ref = ft.Text("", size=12, font_family="Poppins Medium", italic=True, color="grey")
        self.edit_table = ft.ListView(expand=True, divider_thickness=1, spacing=10)
        self.edit_num = ft.Text("", size=12, font_family="Poppins Medium", color="white",)
        self.edit_cree_par = ft.Text("", size=12, font_family="Poppins Medium")
        self.edit_montant = ft.TextField(**readonly_field_style, width=140, prefix_icon=ft.icons.PRICE_CHANGE_OUTLINED, label="Montant")
        self.edit_lettres = ft.TextField(**readonly_field_style, width=700, prefix_icon=ft.icons.LABEL_OUTLINED, label="Montant en lettres")
        self.edit_objet = ft.TextField(**field_style, width=700, label="Objet du devis", prefix_icon=ft.icons.MAIL_OUTLINED)
        self.edit_notabene = ft.TextField(**field_style, width=700, label="NB", prefix_icon=ft.icons.NOTES_OUTLINED)
        self.edit_paiement = ft.TextField(**numbers_field_style, width=110, label="Paiement", prefix_icon=ft.icons.WATCH_LATER_OUTLINED)
        self.edit_validite = ft.TextField(**numbers_field_style, width=120, label="validité", prefix_icon=ft.icons.EVENT_AVAILABLE)
        self.edit_point_liv = ft.TextField(**field_style, width=200, label="pt. livraison", prefix_icon=ft.icons.LOCATION_ON_OUTLINED)
        self.edit_delai = ft.TextField(**field_style, width=150, label="Délai livraison", prefix_icon=ft.icons.TIMELAPSE_OUTLINED)
        self.bt_dim_edit_rem = MiniCtButton(ft.icons.KEYBOARD_ARROW_LEFT_OUTLINED, "Réduire", None,
                                           self.put_edit_remise_down)
        self.bt_add_edit_rem = MiniCtButton(ft.icons.KEYBOARD_ARROW_RIGHT_OUTLINED, "AUgmenter", None,
                                           self.put_edit_remise_up)
        self.edit_remise = ft.TextField(**readonly_field_style, width=60, label="remise", value="0")
        self.bt_valid_modif = AnyButton(FIRST_COLOR, "check", "Valider Modifications", "white", 230, self.update_devis)
        self.bt_print_options = AnyButton(FIRST_COLOR, "print_outlined", "Imprimer devis", "white", 230, self.open_impression_window)
        self.edit_window = ft.Card(
            elevation=20, surface_tint_color="#f0f0f6", width=900, height=650,
            clip_behavior=ft.ClipBehavior.ANTI_ALIAS, shadow_color="black",
            scale=ft.transform.Scale(0), expand=True,
            animate_scale=ft.Animation(300, ft.AnimationCurve.DECELERATE),
            content=ft.Container(
                padding=10, bgcolor="#f0f0f6", expand=True,
                content=ft.Column(
                    controls=[
                        ft.Container(
                            bgcolor="white", padding=ft.padding.only(10, 5, 10, 5), border_radius=16,
                            content=ft.Row(
                                controls=[
                                    ft.Row(
                                        controls=[
                                            ft.Icon(ft.icons.EDIT_OUTLINED, "black"),
                                            ft.Text("Modifier devis".upper(), size=14, font_family="Poppins Bold")
                                        ]
                                    ),
                                    ft.IconButton("close", FIRST_COLOR, scale=0.6, bgcolor="#f2f2f2",
                                                  on_click=self.close_edit_window)
                                ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
                            )
                        ),
                        ft.Container(
                            bgcolor="white", padding=20, border_radius=16, expand=True,
                            content=ft.Column(
                                controls=[
                                    ft.Row(
                                        controls=[
                                            ft.Row(
                                                controls=[
                                                    ft.Row(
                                                        controls=[
                                                            ft.Text("Devis N° ".upper(), size=12,
                                                                    font_family="Poppins Medium"),
                                                            ft.Container(
                                                                padding=ft.padding.only(10, 3, 10, 3), border_radius=10,
                                                                bgcolor=SECOND_COLOR, content=ft.Row(
                                                                    controls=[self.edit_num,], alignment=ft.MainAxisAlignment.CENTER
                                                                )
                                                            ),
                                                        ], spacing=5
                                                    ),
                                                    ft.Row(
                                                        controls=[ft.Text("Créé par ".upper(), size=12, font_family="Poppins Medium"),
                                                            ft.Container(
                                                                padding=ft.padding.only(10, 3, 10, 3), border_radius=10,
                                                                bgcolor="#f2f2f2", content=ft.Row(
                                                                    controls=[self.edit_cree_par],
                                                                    alignment=ft.MainAxisAlignment.CENTER
                                                                )
                                                            )
                                                        ], spacing=5
                                                    ),
                                                    self.edit_bt_facture
                                                ], spacing=30
                                            ),
                                            self.edit_nb_ref
                                        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
                                    ),
                                    ft.Stack(
                                        controls=[
                                            ft.Container(
                                                height=240, padding=ft.padding.only(10, 3, 10, 3),
                                                border_radius=16,
                                                expand=True,
                                                border=ft.border.all(1, "grey"), content=self.edit_table
                                            ),
                                            ft.FloatingActionButton(
                                                bgcolor="black", on_click=self.open_add_ref_window,
                                                content=ft.Row([ft.Icon(ft.icons.ADD, color="white")],
                                                               alignment=ft.MainAxisAlignment.CENTER),
                                                bottom=10, scale=0.8, right=10, tooltip="Ajouter référence", opacity=0.95
                                            )
                                        ], alignment=ft.alignment.bottom_right
                                    ),
                                    self.edit_objet,
                                    ft.Row([self.edit_paiement, self.edit_delai, self.edit_point_liv,
                                            self.edit_validite, ft.Row([self.bt_dim_edit_rem,  self.edit_remise, self.bt_add_edit_rem], spacing=0)]),
                                    self.edit_notabene,
                                    ft.Row([self.edit_montant, self.edit_lettres]),
                                    ft.Row([self.bt_valid_modif, self.bt_print_options], spacing=20)
                                ]
                            )
                        )
                    ]
                )
            )
        )

        # Ajouter reference window ...
        self.add_select_article = ft.TextField(**search_field_style, width=300, on_change=self.filter_selection, prefix_icon="search")
        self.table_edit_ref = ft.ListView(expand=True, divider_thickness=1, spacing=10)
        self.edit_ref_window = ft.Card(
            elevation=20, surface_tint_color="#f0f0f6", width=800, height=550,
            clip_behavior=ft.ClipBehavior.ANTI_ALIAS, shadow_color="black",
            scale=ft.transform.Scale(0), expand=True,
            animate_scale=ft.Animation(300, ft.AnimationCurve.DECELERATE),
            content=ft.Container(
                padding=10, bgcolor="#f0f0f6", expand=True,
                content=ft.Container(
                    expand=True, padding=10, border_radius=16, bgcolor="white",
                    content=ft.Column(
                        expand=True,
                        controls=[
                            ft.Column(
                                controls=[
                                    ft.Text("Ajouter article".upper(), size=14, font_family="Poppins Bold"),
                                    ft.Divider(height=1, thickness=1),
                                ], spacing=0
                            ),
                            self.add_select_article,
                            ft.Container(
                                border_radius=12, border=ft.border.all(1, "grey"), expand=True,
                                padding=ft.padding.only(10, 5, 10, 5),
                                content=self.table_edit_ref
                            ),
                            ft.Row(
                                [AnyButton(FIRST_COLOR, "close", "Fermer", "white", 170, self.close_add_ref_window)],
                                alignment=ft.MainAxisAlignment.END
                            )
                        ]
                    )
                )
            )
        )

        # facture window ...
        self.fac_num_devis = ft.Text("", size=12, font_family="Poppins Medium", color="white")
        self.bc_client = ft.TextField(**field_style, label="BC client", width=270, prefix_icon=ft.icons.CONTENT_PASTE_OUTLINED)
        self.ov_client = ft.TextField(**field_style, label="OV", width=270, prefix_icon=ft.icons.COPY_ALL_OUTLINED)
        self.facture_window = ft.Card(
            elevation=20, surface_tint_color="#f0f0f6", width=300, height=300,
            clip_behavior=ft.ClipBehavior.ANTI_ALIAS, shadow_color="black",
            scale=ft.transform.Scale(0), expand=True,
            animate_scale=ft.Animation(300, ft.AnimationCurve.DECELERATE),
            content=ft.Container(
                padding=10, bgcolor="#f0f0f6",
                content=ft.Container(
                    padding=10, border_radius=16, bgcolor="white",
                    content=ft.Column(
                        controls=[
                            ft.Column(
                                controls=[
                                    ft.Row(
                                        controls=[
                                            ft.Text("Facturer devis".upper(), size=14, font_family="Poppins Bold"),
                                            ft.IconButton("close", FIRST_COLOR, scale=0.6, bgcolor="#f2f2f2", on_click=self.close_facture_window)
                                        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
                                    ),
                                    ft.Divider(height=1, thickness=1),
                                    ft.Divider(height=1, color=ft.colors.TRANSPARENT),
                                    ft.Container(
                                        padding=ft.padding.only(10, 3, 10, 3), bgcolor=SECOND_COLOR, border_radius=10,
                                        content=ft.Row([self.fac_num_devis], alignment=ft.MainAxisAlignment.CENTER)
                                    ),
                                    self.bc_client, self.ov_client
                                ]
                            ),
                            AnyButton(FIRST_COLOR, "check", "Valider facturation", "white", 270, self.facturer_devis)
                        ], horizontal_alignment=ft.CrossAxisAlignment.CENTER
                    )
                )
            )
        )

        # options d'impression window
        self.regime = ft.RadioGroup(
            content=ft.Row(
                controls=[
                    ft.Radio(
                        **radio_style, value="S", label="Simplifié".upper()
                    ),
                    ft.Radio(
                        **radio_style, value="R", label="Réel".upper()
                    )
                ]
            )
        )
        self.banque = ft.RadioGroup(
            content=ft.Row(
                controls=[
                    ft.Radio(
                        **radio_style, value="CCA", label="CCA".upper()
                    ),
                    ft.Radio(
                        **radio_style, value="AFRILAND", label="AFRILAND".upper()
                    )
                ]
            )
        )
        self.tva = ft.Checkbox(
            label_style=ft.TextStyle(size=12, font_family="Poppins Medium"), active_color="white", check_color=FIRST_COLOR,
            label="TVA", label_position=ft.LabelPosition.RIGHT
        )
        self.ir = ft.Checkbox(
            label_style=ft.TextStyle(size=12, font_family="Poppins Medium"), active_color="white", check_color=FIRST_COLOR,
            label="IR", label_position=ft.LabelPosition.RIGHT
        )
        self.download_button = AnyButton(
            FIRST_COLOR, "edit", "Télécharger fichier", "white", 200, None
        )
        self.download_button.visible = False
        self.impression_window = ft.Card(
            elevation=20, surface_tint_color="#f0f0f6", width=350, height=500,
            clip_behavior=ft.ClipBehavior.ANTI_ALIAS, shadow_color="black",
            scale=ft.transform.Scale(0), expand=True,
            animate_scale=ft.Animation(300, ft.AnimationCurve.DECELERATE),
            content=ft.Container(
                padding=10, bgcolor="#f0f0f6",
                content=ft.Container(
                    padding=10, border_radius=16, bgcolor="white",
                    content=ft.Column(
                        controls=[
                            ft.Column(
                                controls=[
                                    ft.Row(
                                        controls=[
                                            ft.Text("Options d'impressions".upper(), size=14, font_family="Poppins Bold"),
                                            ft.IconButton("close", FIRST_COLOR, scale=0.6, bgcolor="#f2f2f2", on_click=self.close_impression_window)
                                        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
                                    ),
                                    ft.Divider(height=1, thickness=1),
                                ],spacing=0
                            ),
                            ft.Divider(height=3, color=ft.colors.TRANSPARENT),
                            ft.Column(
                                controls=[
                                    ft.Text("choix du régime".upper(), size=11,font_family="Poppins Bold"),
                                    ft.Divider(height=1, thickness=1),
                                ],spacing=0
                            ),
                            self.regime,
                            ft.Divider(height=3, color=ft.colors.TRANSPARENT),
                            ft.Column(
                                controls=[
                                    ft.Text("TVA et IR".upper(), size=11, font_family="Poppins Bold"),
                                    ft.Divider(height=1, thickness=1),
                                ], spacing=0
                            ),
                            ft.Row([self.tva, self.ir]),
                            ft.Divider(height=1, color=ft.colors.TRANSPARENT),
                            ft.Column(
                                controls=[
                                    ft.Text("Choix de la banque".upper(), size=11, font_family="Poppins Bold"),
                                    ft.Divider(height=1, thickness=1),
                                ], spacing=0
                            ),
                            self.banque,
                            ft.Divider(height=1, color=ft.colors.TRANSPARENT),
                            AnyButton(
                                FIRST_COLOR, ft.icons.LOCAL_PRINTSHOP_OUTLINED, "Imprimer", "white", 280,
                                self.imprimer_devis
                            ),
                            self.download_button
                        ], horizontal_alignment=ft.CrossAxisAlignment.CENTER
                    )
                )
            )
        )

        self.content = ft.Stack(
            controls=[
                self.main_window, self.new_window, self.new_add_ref_window, self.edit_window,
                self.edit_ref_window, self.facture_window, self.impression_window
            ], alignment=ft.alignment.center
        )
        self.load_datas()
        self.load_lists()

    def load_lists(self):
        datas = be.all_references()

        # on charge liste des articles pour les creations de devis ...
        for widget in self.new_table_add_ref.controls[:]:
            self.new_table_add_ref.controls.remove(widget)

        for data in datas:
            self.new_table_add_ref.controls.append(
                OneArticle(self, data["reference"], data["designation"], data["prix"])
            )

        # on charge liste des articles pour les modifications de devis ...
        for widget in self.table_edit_ref.controls[:]:
            self.table_edit_ref.controls.remove(widget)

        for data in datas:
            self.table_edit_ref.controls.append(
                AddOneArticle(self, data["reference"], data["designation"], data["prix"])
            )

        # on charge la liste des clients ...
        clients = be.all_clients()
        for client in clients:
            self.new_client.options.append(
                ft.dropdown.Option(client["nom"])
            )

    def load_datas(self):
        datas = be.all_devis()
        self.results.value = f"{len(datas)} Résultat(s)"

        for row in self.table.rows[:]:
            self.table.rows.remove(row)

        for data in datas:
            if data["statut"].lower() != "non facturé":
                icone = ft.icons.CREDIT_SCORE_ROUNDED
                couleur = ft.colors.DEEP_ORANGE
                bill_button = CtButton(None, ft.colors.BLACK45, "", data, None)
            else:
                icone = None
                couleur = None
                bill_button = CtButton(ft.icons.ADD_CARD, ft.colors.BLACK45, "Facturer devis", data, self.open_facture_window)

            self.table.rows.append(
                ft.DataRow(
                    cells=[
                        ft.DataCell(ft.Icon(icone, couleur, 20)),
                        ft.DataCell(ft.Text(data["date"])),
                        ft.DataCell(ft.Text(data["numero"])),
                        ft.DataCell(ft.Text(data["client"])),
                        ft.DataCell(ft.Text(ajout_separateur(data["montant"]))),
                        ft.DataCell(
                            ft.Row(
                                controls=[
                                    CtButton("edit_outlined", ft.colors.BLUE_300,"Modifier", data, self.open_edit_window),
                                    bill_button
                                ], alignment=ft.MainAxisAlignment.END, spacing=0,
                            )
                        )
                    ]
                )
            )
        datas = be.all_references()

    def filter_datas(self, e):
        datas = be.all_devis()
        search = self.search.value if self.search.value is not None else ""
        filtered_datas = list(filter(lambda x: search in x['client'] or search in x["numero"], datas))

        self.results.value = f"{len(filtered_datas)} Résultat(s)"
        self.results.update()

        for row in self.table.rows[:]:
            self.table.rows.remove(row)

        for data in filtered_datas:
            if data["statut"].lower() != "non facturé":
                icone = ft.icons.CREDIT_SCORE_ROUNDED
                couleur = ft.colors.DEEP_ORANGE
                bill_button = CtButton(None, ft.colors.BLACK45, "", data, None)
            else:
                icone = None
                couleur = None
                bill_button = CtButton(ft.icons.ADD_CARD, ft.colors.BLACK45, "Facturer devis", data, self.open_facture_window)

            self.table.rows.append(
                ft.DataRow(
                    cells=[
                        ft.DataCell(ft.Icon(icone, couleur, 20)),
                        ft.DataCell(ft.Text(data["date"])),
                        ft.DataCell(ft.Text(data["numero"])),
                        ft.DataCell(ft.Text(data["client"])),
                        ft.DataCell(ft.Text(ajout_separateur(data["montant"]))),
                        ft.DataCell(
                            ft.Row(
                                controls=[
                                    CtButton("edit_outlined", ft.colors.BLUE_300, "Modifier", data, self.open_edit_window),
                                    bill_button
                                ], alignment=ft.MainAxisAlignment.END, spacing=0,
                            )
                        )
                    ]
                )
            )

        self.table.update()

    def filter_selection(self, e):
        selected = self.add_select_article.value if self.add_select_article.value is not None else ""
        datas = be.all_references()

        for widget in self.table_edit_ref.controls[:]:
            self.table_edit_ref.controls.remove(widget)

        filtered_datas = list(filter(lambda x: selected in x['designation'], datas))

        for data in filtered_datas:
            self.table_edit_ref.controls.append(
                OneArticle(self, data["reference"], data["designation"], data["prix"])
            )
        self.table_edit_ref.update()

    def filter_selection_new(self, e):
        selected = self.new_add_select_article.value if self.new_add_select_article.value is not None else ""
        datas = be.all_references()

        for widget in self.new_table_add_ref.controls[:]:
            self.new_table_add_ref.controls.remove(widget)

        filtered_datas = list(filter(lambda x: selected in x['designation'], datas))

        for data in filtered_datas:
            self.new_table_add_ref.controls.append(
                OneArticle(self, data["reference"], data["designation"], data["prix"])
            )
        self.new_table_add_ref.update()

    def open_new_window(self, e):
        self.new_window.scale = 1
        self.new_window.update()
        self.main_window.opacity = 0.3
        self.main_window.disabled = True
        self.main_window.update()

    def open_edit_window(self, e):
        datas = be.select_one_devis(e.control.data["numero"])
        self.edit_cree_par.value = f"{datas['cree_par']}".upper()
        self.edit_objet.value = f"{datas["objet"]}"
        self.edit_notabene.value = f"{datas["note_bene"]}"
        self.edit_paiement.value = f"{datas["paiement"]}"
        self.edit_validite.value = f"{datas["validite"]}"
        self.edit_point_liv.value = f"{datas["point_liv"]}"
        self.edit_delai.value = f"{datas["delai"]}"
        self.edit_remise.value = f"{datas["remise"]}"
        self.edit_montant.value = f"{datas["montant"]}"
        self.edit_num.value = f"{datas["numero"]}"
        self.edit_lettres.value = ecrire_en_lettres(int(datas["montant"]))

        for widget in (
            self.edit_objet, self.edit_delai, self.edit_validite, self.edit_remise, self.edit_montant,
            self.edit_paiement, self.edit_point_liv, self.edit_notabene, self.edit_num, self.edit_lettres,
            self.edit_cree_par
        ):
            widget.update()

        details = be.find_devis_details(e.control.data["numero"])
        self.edit_nb_ref.value = f"{len(details)} ligne(s)"
        self.edit_nb_ref.update()

        for widget in self.edit_table.controls[:]:
            self.edit_table.controls.remove(widget)

        for detail in details:
            designation = be.search_designation(detail["reference"])
            self.edit_table.controls.append(
                EditOneArticle(self, detail["reference"], designation, detail["qte"], detail["prix"])
            )

        if e.control.data["statut"].lower() == "facturé":
            self.edit_bt_facture.content = ft.Row(
                controls=[
                    ft.Icon(ft.icons.CHECK, color="white", size=16),
                    ft.Text("facturé".upper(), size=12, font_family="Poppins Medium", color="white",)
                ], spacing=5
            )
            self.edit_bt_facture.bgcolor = ft.colors.GREEN
            self.edit_bt_facture.update()
            self.bt_valid_modif.visible = False
            self.bt_valid_modif.update()

        else:
            self.edit_bt_facture.content = ft.Row(
                controls=[
                    ft.Icon(ft.icons.INDETERMINATE_CHECK_BOX_OUTLINED, color="white", size=16),
                    ft.Text("non facturé".upper(), size=12, font_family="Poppins Medium", color="white",)
                ], spacing=5
            )
            self.edit_bt_facture.bgcolor = "red"
            self.edit_bt_facture.update()
            self.bt_valid_modif.visible = True
            self.bt_valid_modif.update()

        self.edit_table.update()
        self.edit_window.scale = 1
        self.edit_window.update()
        self.main_window.opacity = 0.3
        self.main_window.disabled = True
        self.main_window.update()

    def close_new_window(self, e):
        self.new_window.scale = 0
        self.new_window.update()
        self.main_window.opacity = 1
        self.main_window.disabled = False
        self.main_window.update()

    def close_edit_window(self, e):
        self.edit_window.scale = 0
        self.edit_window.update()
        self.main_window.opacity = 1
        self.main_window.disabled = False
        self.main_window.update()

    def close_add_ref_window(self, e):
        self.edit_ref_window.scale = 0
        self.edit_ref_window.update()

    def open_add_ref_window(self, e):
        self.edit_ref_window.scale = 1
        self.edit_ref_window.update()

    def close_new_add_ref_window(self, e):
        self.new_add_ref_window.scale = 0
        self.new_add_ref_window.update()

    def open_new_add_ref_window(self, e):
        self.new_add_ref_window.scale = 1
        self.new_add_ref_window.update()

    # A chaque fois que le client change
    def on_change_new_client(self, e):
        id_client = be.id_client_by_name(self.new_client.value)
        numero_de_devis = be.find_devis_num(id_client)
        self.new_num.value = f"{numero_de_devis}"
        self.new_num.update()

    # A chaque changement de remise
    def put_new_remise_up(self, e):
        anc_remise = int(self.new_remise.value)

        next_remise = anc_remise + 1
        self.new_remise.value = f"{next_remise}"
        self.new_remise.update()

        total = 0
        for widget in self.new_table.controls[:]:
            total += (int(widget.art_qte.value) * int(widget.art_prix.value))

        montant_remise = total - (total * next_remise / 100)
        self.new_montant.value = f"{int(montant_remise)}"
        self.new_lettres.value = ecrire_en_lettres(int(montant_remise))

        for widget in (self.new_lettres, self.new_montant):
            widget.update()

    def put_new_remise_down(self, e):
        anc_remise = int(self.new_remise.value)

        if anc_remise == 0:
            next_remise = 0
        else:
            next_remise = anc_remise - 1

        self.new_remise.value = f"{next_remise}"
        self.new_remise.update()
        total = 0
        for widget in self.new_table.controls[:]:
            total += (int(widget.art_qte.value) * int(widget.art_prix.value))

        montant_remise = total - (total * next_remise / 100)
        self.new_montant.value = f"{int(montant_remise)}"
        self.new_lettres.value = ecrire_en_lettres(int(montant_remise))

        for widget in (self.new_lettres, self.new_montant):
            widget.update()

    # A chaque changement de remise dans la modificatyion du devis
    def put_edit_remise_up(self, e):
        anc_remise = int(self.edit_remise.value)
        next_remise = anc_remise + 1

        self.edit_remise.value = f"{next_remise}"
        self.edit_remise.update()
        total = 0
        for widget in self.edit_table.controls[:]:
            total += (int(widget.art_qte.value) * int(widget.art_prix.value))

        montant_remise = total - (total * next_remise / 100)
        self.edit_montant.value = f"{int(montant_remise)}"
        self.edit_lettres.value = ecrire_en_lettres(int(montant_remise))

        for widget in (self.edit_lettres, self.edit_montant):
            widget.update()

    def put_edit_remise_down(self, e):
        anc_remise = int(self.edit_remise.value)

        if anc_remise == 0:
            next_remise = 0
        else:
            next_remise = anc_remise - 1

        self.edit_remise.value = f"{next_remise}"
        self.edit_remise.update()
        total = 0
        for widget in self.edit_table.controls[:]:
            total += (int(widget.art_qte.value)*int(widget.art_prix.value))

        montant_remise = total - (total*next_remise/100)
        self.edit_montant.value = f"{int(montant_remise)}"
        self.edit_lettres.value = ecrire_en_lettres(int(montant_remise))

        for widget in (self.edit_lettres, self.edit_montant):
            widget.update()

    # Création d'un nouveau devis
    def create_new_devis(self, e):

        if self.new_client.value == "" or self.new_client.value is None:
            self.cp.box.title.value = "Erreur"
            self.cp.box.content.value = "Cleint obligatoire"
            self.cp.box.open = True
            self.cp.box.update()
        else:
            for row in self.new_table.controls[:]:
                be.add_devis_details(self.new_num.value, row.ref, int(row.art_qte.value), int(row.art_prix.value))

            new_date = datetime.date.today()
            id_client = be.id_client_by_name(self.new_client.value)
            objet = self.new_objet.value if self.new_objet.value is not None else ""
            note_bene = self.new_notabene.value if self.new_notabene.value is not None else ""
            remise = int(self.new_remise.value) if self.new_remise.value is not None else 0
            mt_lettres = ""
            note_bene = self.new_notabene.value if self.new_notabene.value is not None else ""
            delai = self.new_delai.value if self.new_delai.value is not None else ""
            point_liv = self.new_point_liv.value if self.new_point_liv.value is not None else ""
            validite = self.new_validite.value if self.new_validite.value is not None else 0
            paiement = int(self.new_paiement.value) if self.new_paiement.value is not None else 0

            be.add_devis(
                self.new_num.value, str(new_date), id_client, int(self.new_montant.value),
                objet, remise, mt_lettres, note_bene, delai, point_liv, validite, paiement,
                self.cp.user_infos['username']
            )

            self.cp.box.title.value = "Validé"
            self.cp.box.content.value = f"Devis N° {self.new_num.value} créé"
            self.cp.box.open = True
            self.cp.box.update()

            for widget in self.new_table.controls[:]:
                self.new_table.controls.remove(widget)

            self.new_table.update()

            self.load_datas()
            self.table.update()
            self.results.update()

            for widget in (self.new_client, self.new_num, self.new_objet, self.new_notabene, self.new_delai, self.new_point_liv, self.new_validite):
                widget.value = None
                widget.update()

            for widget in (self.new_montant, self.new_remise):
                widget.value = "0"
                widget.update()

            self.new_lettres.value = ecrire_en_lettres(int(self.new_montant.value))
            self.new_lettres.update()

    def update_devis(self, e):
        num_devis = self.edit_num.value

        # on met à jour la table devis
        mt = self.edit_montant.value

        be.update_devis(
            mt, int(self.edit_remise.value), self.edit_notabene.value, self.edit_delai.value, self.edit_point_liv.value,
            int(self.edit_validite.value), int(self.edit_paiement.value), num_devis
        )

        # on supprime les details devis précéédents
        be.delete_devis_details(num_devis)

        # on recrée les nouveaux détails
        for widget in self.edit_table.controls[:]:
            be.add_devis_details(num_devis, widget.ref, int(widget.art_qte.value), int(widget.art_prix.value))

        self.edit_window.scale = 0
        self.edit_window.update()

        self.load_datas()
        self.table.update()
        self.results.update()

        self.cp.box.title.value = "Validé"
        self.cp.box.content.value = "devis mis à jour".capitalize()
        self.cp.box.open = True
        self.cp.box.update()

    def open_facture_window(self, e):
        self.fac_num_devis.value = e.control.data["numero"]
        self.fac_num_devis.update()
        self.facture_window.scale = 1
        self.facture_window.update()
        self.main_window.opacity = 0.3
        self.main_window.disabled = True
        self.main_window.update()

    def close_facture_window(self, e):
        self.bc_client.value = None
        self.bc_client.update()
        self.ov_client.value = None
        self.ov_client.update()
        self.facture_window.scale = 0
        self.facture_window.update()
        self.main_window.opacity = 1
        self.main_window.disabled = False
        self.main_window.update()

    # Factuer devis
    def facturer_devis(self, e):
        details = be.find_devis_details(self.fac_num_devis.value)

        # On vérifie les si le stock est disponible pour chaque produit de natuire stock ...
        count_nc = 0
        for row in details:
            if be.find_nature_ref(row["reference"]) == "stock":
                ancien_stock = be.find_stock_ref(row["reference"])

                if ancien_stock < row["qte"]:
                    count_nc += 1

        if count_nc == 0:
            info_facture = be.show_info_devis(self.fac_num_devis.value)
            numero_facture = be.find_facture_num(info_facture["client"])

            # table facture
            be.add_facture(
                numero_facture, info_facture["client"], info_facture["montant"], info_facture["objet"], info_facture["remise"], "",
                self.fac_num_devis.value, self.bc_client.value, self.ov_client.value, info_facture["paiement"])

            # Table details facture
            for row in details:
                be.add_details_facture(numero_facture, row["reference"], row["qte"], row["prix"])
                # Mise à jour du stock
                if be.find_nature_ref(row["reference"]) == "stock":
                    ancien_stock = be.find_stock_ref(row["reference"])
                    nouveau_stock = ancien_stock - row["qte"]
                    be.update_stock(nouveau_stock, row["reference"])
                    be.add_historique(row[2], "S", numero_facture, ancien_stock, row["qte"], nouveau_stock)

            # mise à jour du statut du devis
            be.maj_statut_devis(self.fac_num_devis.value)

            # remplir les bordereaux de livraison
            initiales_client = be.search_initiales(info_facture["client"])
            numero_bordereau = be.find_bordereau_num(initiales_client)
            be.add_bordereau(numero_bordereau, numero_facture, self.bc_client.value)

            for row in details:
                be.add_bordereau_details(numero_bordereau, row["reference"], row["qte"], row["prix"])

            self.bc_client.value = None
            self.bc_client.update()
            self.ov_client.value = None
            self.ov_client.update()
            self.facture_window.scale = 0
            self.facture_window.update()
            self.load_datas()
            self.table.update()
            self.results.update()

            self.cp.box.title.value = "Validé"
            self.cp.box.content.value = f"facture N° {numero_facture} générée avec succés"
            self.cp.box.open = True
            self.cp.box.update()

        else:
            self.cp.box.title.value = "Erreur"
            self.cp.box.content.value = f"Veuiullez vérifier les quantités en stock"
            self.cp.box.open = True
            self.cp.box.update()

    def open_impression_window(self, e):
        self.impression_window.scale = 1
        self.impression_window.update()

    def close_impression_window(self, e):
        self.impression_window.scale = 0
        self.impression_window.update()

    def imprimer_devis(self, e: ft.FilePickerResultEvent):
        regime = self.regime.value
        tva = self.tva.value
        ir = self.ir.value

        # Cas du régime simplifié
        if self.regime.value == "S":

            # Si la TVA est Active
            if tva is True:
                # Erreur
                self.cp.box.title.value = "Erreur"
                self.cp.box.content.value = "Pas de TVA dans le régime simplifié"
                self.cp.box.open = True
                self.cp.box.update()

            # Si la TVA n'est pas active
            else:
                # Si l'IR est active
                if ir is True:
                    # generer le document word
                    total_prix = 0
                    def generate_word_doc():
                        doc = Document()

                        # Ajouter une image dans l'en-tête
                        def header_and_footer():
                            section = doc.sections[0]
                            header = section.header
                            header_paragraph = header.paragraphs[0]
                            footer = section.footer
                            footer_paragraph = footer.paragraphs[0]
                            # Ajouter l'image dans l'en-tête
                            header_paragraph.add_run().add_picture("assets/images/header.jpg", width=Inches(6.5))
                            # Ajouter l'image dans le pied de page
                            footer_paragraph.add_run().add_picture("assets/images/footer.png", width=Inches(6.5))

                        header_and_footer()

                        # Créer un tableau principal avec une seule ligne et deux colonnes pour entete client
                        main_table = doc.add_table(rows=1, cols=2)

                        # Entête infos facture
                        def entete_facture():
                            # Ajouter un tableau dans la première cellule
                            cell1 = main_table.cell(0, 0)
                            num_proforma = self.edit_num.value
                            suivant = "Suivant proforma N°"
                            date = be.show_info_devis(num_proforma)["date"]

                            table1 = cell1.add_table(rows=4, cols=1)  # Tableau avec 4 lignes et 1 colonne
                            cell1_1 = table1.cell(0, 0)
                            paragraph1 = cell1_1.paragraphs[0]
                            run1 = paragraph1.add_run("PROFORMA")
                            paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run1.font.name = "Arial Black"
                            run1.font.size = Pt(16)
                            cell1_1.paragraphs[0].paragraph_format.space_before = 0  # Pas d'espace avant le paragraphe
                            cell1_1.paragraphs[0].paragraph_format.space_after = 0  # Pas d'espace après le paragraphe
                            cell1_1.paragraphs[0].paragraph_format.line_spacing = Pt(25)  # Espacement entre les lignes réduit

                            cell1_2 = table1.cell(1, 0)
                            paragraph2 = cell1_2.paragraphs[0]
                            run2 = paragraph2.add_run(f"{num_proforma}")
                            paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run2.font.name = "calibri"
                            run2.font.size = Pt(12)
                            cell1_2.paragraphs[0].paragraph_format.space_before = 10  # Pas d'espace avant le paragraphe
                            cell1_2.paragraphs[0].paragraph_format.space_after = 10  # Pas d'espace après le paragraphe
                            cell1_2.paragraphs[0].paragraph_format.line_spacing = Pt(15)  # Espacement entre les lignes réduit

                            cell1_3 = table1.cell(2, 0)
                            paragraph3 = cell1_3.paragraphs[0]
                            run3 = paragraph3.add_run(f"Suivant demande du {ecrire_date(date)}")
                            paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run2.font.name = "calibri"
                            run2.font.size = Pt(12)
                            cell1_3.paragraphs[0].paragraph_format.space_before = 10  # Pas d'espace avant le paragraphe
                            cell1_3.paragraphs[0].paragraph_format.space_after = 10  # Pas d'espace après le paragraphe
                            cell1_3.paragraphs[0].paragraph_format.line_spacing = Pt(15)  # Espacement entre les lignes réduit

                        entete_facture()

                        # Entete infos client
                        def entete_infos_client():
                            # Ajouter un tableau dans la première cellule
                            cell1 = main_table.cell(0, 1)
                            client_id = be.show_info_devis(self.edit_num.value)["client"]
                            client = be.infos_clients(client_id)["nom"]
                            contact = be.infos_clients(client_id)["contact"]
                            nui = be.infos_clients(client_id)["NUI"]
                            rc = be.infos_clients(client_id)["RC"]

                            table1 = cell1.add_table(rows=4, cols=1)  # Tableau avec 4 lignes et 1 colonne
                            cell1_1 = table1.cell(0, 0)
                            paragraph1 = cell1_1.paragraphs[0]
                            run1 = paragraph1.add_run(f"Client: {client}")
                            paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run1.font.name = "calibri"
                            run1.font.size = Pt(12)
                            cell1_1.paragraphs[0].paragraph_format.space_before = 0  # Pas d'espace avant le paragraphe
                            cell1_1.paragraphs[0].paragraph_format.space_after = 0  # Pas d'espace après le paragraphe
                            cell1_1.paragraphs[0].paragraph_format.line_spacing = Pt(
                                20)  # Espacement entre les lignes réduit

                            cell1_2 = table1.cell(1, 0)
                            paragraph2 = cell1_2.paragraphs[0]
                            run2 = paragraph2.add_run(f"Contact: {contact}")
                            paragraph2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run2.font.name = "calibri"
                            run2.font.size = Pt(12)
                            cell1_2.paragraphs[0].paragraph_format.space_before = 10  # Pas d'espace avant le paragraphe
                            cell1_2.paragraphs[0].paragraph_format.space_after = 10  # Pas d'espace après le paragraphe
                            cell1_2.paragraphs[0].paragraph_format.line_spacing = Pt(15)

                            cell1_3 = table1.cell(2, 0)
                            paragraph3 = cell1_3.paragraphs[0]
                            run3 = paragraph3.add_run(f"NUI: {nui}")
                            paragraph3.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run3.font.name = "calibri"
                            run3.font.size = Pt(12)
                            cell1_3.paragraphs[0].paragraph_format.space_before = 10  # Pas d'espace avant le paragraphe
                            cell1_3.paragraphs[0].paragraph_format.space_after = 10  # Pas d'espace après le paragraphe
                            cell1_3.paragraphs[0].paragraph_format.line_spacing = Pt(15)

                            cell1_4 = table1.cell(3, 0)
                            paragraph4 = cell1_4.paragraphs[0]
                            run4 = paragraph4.add_run(f"RC: {rc}")
                            paragraph4.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run4.font.name = "calibri"
                            run4.font.size = Pt(12)
                            cell1_4.paragraphs[0].paragraph_format.space_before = 10  # Pas d'espace avant le paragraphe
                            cell1_4.paragraphs[0].paragraph_format.space_after = 10  # Pas d'espace après le paragraphe
                            cell1_4.paragraphs[0].paragraph_format.line_spacing = Pt(15)

                        entete_infos_client()

                        # Objet
                        def draw_simple_paragraph(
                                text: str, alignment, before: int, after: int, font_size: int,
                                is_italic: bool, is_bold: bool):
                            details_pg = doc.add_paragraph()
                            details_pg.alignment = alignment
                            details_pg_format = details_pg.paragraph_format
                            details_pg_format.space_before = Pt(before)  # Espace avant le paragraphe
                            details_pg_format.space_after = Pt(after)
                            details_pg_format.line_spacing = 1
                            run_details_header = details_pg.add_run(text)  # Run 1 Separation
                            run_details_header.font.name = "calibri"
                            run_details_header.font.size = Pt(font_size)
                            run_details_header.italic = is_italic
                            run_details_header.bold = is_bold

                        # Objet
                        objet = be.show_info_devis(self.edit_num.value)['objet']
                        if objet == "" or objet is None:
                            pass
                        else:
                            draw_simple_paragraph(
                                f"Objet: {objet}", WD_PARAGRAPH_ALIGNMENT.LEFT,
                                10, 10, 12, False, False
                            )

                        # References
                        def draw_details_devis():
                            details = be.find_devis_details(self.edit_num.value)
                            longueur = len(details) + 1

                            # Créer un tableau avec 3 lignes et 6 colonnes
                            table = doc.add_table(rows=longueur, cols=6)

                            # Appliquer des bordures au tableau entier
                            # Ajouter des bordures noires à chaque cellule
                            def set_cell_border(cell):
                                """ Ajoute une bordure noire autour d'une cellule """
                                tc_pr = cell._element.get_or_add_tcPr()
                                borders = OxmlElement('w:tcBorders')

                                for border_name in ['top', 'left', 'bottom', 'right']:
                                    border = OxmlElement(f'w:{border_name}')
                                    border.set(qn('w:val'), 'single')  # Bordure simple
                                    border.set(qn('w:sz'),
                                               '8')  # Taille de la bordure (plus épais pour meilleure visibilité)
                                    border.set(qn('w:space'), '0')
                                    border.set(qn('w:color'), '000000')  # Noir
                                    borders.append(border)

                                tc_pr.append(borders)

                            # Appliquer les bordures à toutes les cellules
                            for row in table.rows:
                                for cell in row.cells:
                                    set_cell_border(cell)


                            # Définir la largeur spécifique des colonnes (en centimètres)
                            column_widths = [Cm(1), Cm(12), Cm(1.5), Cm(1), Cm(2), Cm(2.5)]  # Largeurs des colonnes

                            # Appliquer les largeurs de colonnes
                            for col_idx, width in enumerate(column_widths):
                                for row in table.rows:
                                    row.cells[col_idx].width = width  # Définir la largeur de la cellule

                            # Ajouter les en-têtes (ligne 0)
                            hdr_cells = table.rows[0].cells
                            headers = ["item", "Désignation", "Qté", "U", "Prix", "total (CFA)"]
                            for i, hdr in enumerate(headers):
                                paragraph = hdr_cells[i].paragraphs[0]
                                run = paragraph.add_run(hdr)
                                run.bold = True
                                run.font.name = "Calibri"
                                run.font.size = Pt(10)

                                # Centrer horizontalement
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                # Centrer verticalement


                            # Ajouter des données dans les 2 autres lignes (lignes 1 et 2) avec des nombres aléatoires
                            for i, row in enumerate(table.rows[1:]):  # i commence à 0 pour la première ligne de données
                                # Indice de la ligne (Ligne 1, Ligne 2, etc.)
                                details = be.find_devis_details(self.edit_num.value)
                                values = [f"{i }", str(details[i]["designation"]), str(details[i]["qte"]),
                                          str(details[i]["unite"]), str(ajout_separateur(details[i]["prix"])),
                                          str(ajout_separateur(details[i]["prix"]*details[i]["qte"]))
                                ]

                                for j, cell in enumerate(row.cells):
                                    paragraph = cell.paragraphs[0]
                                    run = paragraph.add_run(values[j])
                                    run.font.name = "Calibri"
                                    run.font.size = Pt(10)
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrer le texte
                                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER



                            # Appliquer des bordures noires à chaque cellule
                            # for row in table.rows:
                            #     for cell in row.cells:
                            #         cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                            #         borders = cell._element.xpath('.//w:tcBorders')[0]
                            #         for border in borders:
                            #             border.set(qn('w:top'), 'single')
                            #             border.set(qn('w:left'), 'single')
                            #             border.set(qn('w:bottom'), 'single')
                            #             border.set(qn('w:right'), 'single')

                        draw_details_devis()

                        # Montant total
                        mt_total = 0
                        details = be.find_devis_details(self.edit_num.value)
                        for row in details:
                            mt_total += row["qte"]*row["prix"]

                        def draw_montants():
                            draw_simple_paragraph(
                                f"Total: {ajout_separateur(mt_total)}", WD_PARAGRAPH_ALIGNMENT.RIGHT, 10, 1,
                                11, False, False
                            )
                            if int(self.edit_remise.value) == 0:
                                mt_ir = int(mt_total*2.2/100)
                                draw_simple_paragraph(
                                    f"IR: {ajout_separateur(mt_ir)} %",
                                    WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                    11, False, False
                                )
                                mt_nap = int(mt_total - mt_ir)
                                draw_simple_paragraph(
                                    f"NAP: {ajout_separateur(mt_nap)} %",
                                    WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                    11, False, False
                                )
                            else:
                                draw_simple_paragraph(
                                    f"Remise: {self.edit_remise.value} %", WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                    11, False, False
                                )
                                mt_remise = be.show_info_devis(self.edit_num.value)['montant']
                                draw_simple_paragraph(
                                    f"Montant Remisé: {ajout_separateur(mt_remise)} %",
                                    WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                    11, False, False
                                )
                                mt_ir = int(mt_remise*2.2/100)
                                draw_simple_paragraph(
                                    f"IR: {ajout_separateur(mt_ir)} %",
                                    WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                    11, False, False
                                )
                                mt_nap = mt_remise - mt_ir
                                draw_simple_paragraph(
                                    f"NAP: {ajout_separateur(mt_nap)} %",
                                    WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                    11, False, False
                                )

                            draw_simple_paragraph(
                                f"Facture Proforma arrêtée à la somme de:", WD_PARAGRAPH_ALIGNMENT.CENTER, 10, 1,
                                10, True, False
                            )
                            draw_simple_paragraph(
                                f"{ecrire_en_lettres(mt_nap)}".upper(), WD_PARAGRAPH_ALIGNMENT.CENTER, 1, 10,
                                11, False, True
                            )

                            if be.show_info_devis(self.edit_num.value)["note_bene"] == "" or \
                                    be.show_info_devis(self.edit_num.value)["note_bene"] is None:
                                pass
                            else:
                                draw_simple_paragraph(
                                    f"{ecrire_en_lettres(mt_nap)}".upper(), WD_PARAGRAPH_ALIGNMENT.CENTER, 1, 10,
                                    11, False, True
                                )
                                observations = be.show_info_devis(self.edit_num.value)["note_bene"].split(";")
                                for observ in observations:
                                    draw_simple_paragraph(
                                        f"{observ}", WD_PARAGRAPH_ALIGNMENT.LEFT, 20, 10,
                                        11, False, False
                                    )

                        draw_montants()






                        # Enregistrement du fichier
                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        return buffer.getvalue()

                    def upload_to_supabase(file, filename):
                        """Upload un fichier sur Supabase et retourne le lien de téléchargement"""
                        time_stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                        file_path = f"{time_stamp}_{filename}"

                        # 🔹 Upload vers Supabase Storage
                        resp = supabase.storage.from_("devis").upload(
                            file_path,
                            file,  # Fichier en bytes
                            file_options={
                                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
                            # Ajout du MIME type ici
                        )
                        # 🔹 Vérification des erreurs
                        if isinstance(resp, dict) and "error" in resp and resp["error"]:
                            return None, resp["error"]

                        # 🔹 Générer l'URL publique
                        url = supabase.storage.from_("devis").get_public_url(file_path)
                        return url, None

                    file_bytes = generate_word_doc()
                    file_url, error = upload_to_supabase(file_bytes, "mon_devis.docx")

                    if error:
                        self.download_button.disabled = True
                    else:
                        self.download_button.url = file_url
                        self.download_button.visible = True
                        self.download_button.update()


                # Si l'IR n'est pos active
                else:
                    # generer le document word
                    def generate_word_doc():
                        doc = Document()
                        # Ajouter une image dans l'en-tête
                        section = doc.sections[0]
                        header = section.header
                        header_paragraph = header.paragraphs[0]
                        footer = section.footer
                        footer_paragraph = footer.paragraphs[0]

                        # Ajouter l'image dans l'en-tête
                        header_paragraph.add_run().add_picture("assets/images/header.jpg", width=Inches(6.5))
                        # Ajouter l'image dans le pied de page
                        footer_paragraph.add_run().add_picture("assets/images/footer.png", width=Inches(6.5))

                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        return buffer.getvalue()

                    def upload_to_supabase(file, filename):
                        """Upload un fichier sur Supabase et retourne le lien de téléchargement"""
                        time_stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                        file_path = f"{time_stamp}_{filename}"

                        # 🔹 Upload vers Supabase Storage
                        resp = supabase.storage.from_("devis").upload(
                            file_path,
                            file,  # Fichier en bytes
                            file_options={
                                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
                            # Ajout du MIME type ici
                        )
                        # 🔹 Vérification des erreurs
                        if isinstance(resp, dict) and "error" in resp and resp["error"]:
                            return None, resp["error"]

                        # 🔹 Générer l'URL publique
                        url = supabase.storage.from_("devis").get_public_url(file_path)
                        return url, None

                    file_bytes = generate_word_doc()
                    file_url, error = upload_to_supabase(file_bytes, "mon_devis.docx")

                    if error:
                        self.download_button.disabled = True
                    else:
                        self.download_button.url = file_url
                        self.download_button.visible = True
                        self.download_button.update()

        # Cas du régime réel
        else:
            # Si la TVA est Active
            if tva is True:
                if ir is True:
                    # generer le document word
                    def generate_word_doc():
                        doc = Document()
                        # Ajouter une image dans l'en-tête
                        section = doc.sections[0]
                        header = section.header
                        header_paragraph = header.paragraphs[0]
                        footer = section.footer
                        footer_paragraph = footer.paragraphs[0]

                        # Ajouter l'image dans l'en-tête
                        header_paragraph.add_run().add_picture("assets/images/header.jpg", width=Inches(6.5))
                        # Ajouter l'image dans le pied de page
                        footer_paragraph.add_run().add_picture("assets/images/footer.png", width=Inches(6.5))

                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        return buffer.getvalue()

                    def upload_to_supabase(file, filename):
                        """Upload un fichier sur Supabase et retourne le lien de téléchargement"""
                        time_stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                        file_path = f"{time_stamp}_{filename}"

                        # 🔹 Upload vers Supabase Storage
                        resp = supabase.storage.from_("devis").upload(
                            file_path,
                            file,  # Fichier en bytes
                            file_options={
                                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
                            # Ajout du MIME type ici
                        )
                        # 🔹 Vérification des erreurs
                        if isinstance(resp, dict) and "error" in resp and resp["error"]:
                            return None, resp["error"]

                        # 🔹 Générer l'URL publique
                        url = supabase.storage.from_("devis").get_public_url(file_path)
                        return url, None

                    file_bytes = generate_word_doc()
                    file_url, error = upload_to_supabase(file_bytes, "mon_devis.docx")

                    if error:
                        self.download_button.disabled = True
                    else:
                        self.download_button.url = file_url
                        self.download_button.visible = True
                        self.download_button.update()

                # Si l'IR n'est pos active
                else:
                    # generer le document word
                    def generate_word_doc():
                        doc = Document()
                        # Ajouter une image dans l'en-tête
                        section = doc.sections[0]
                        header = section.header
                        header_paragraph = header.paragraphs[0]
                        footer = section.footer
                        footer_paragraph = footer.paragraphs[0]

                        # Ajouter l'image dans l'en-tête
                        header_paragraph.add_run().add_picture("assets/images/header.jpg", width=Inches(6.5))
                        # Ajouter l'image dans le pied de page
                        footer_paragraph.add_run().add_picture("assets/images/footer.png", width=Inches(6.5))

                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        return buffer.getvalue()

                    def upload_to_supabase(file, filename):
                        """Upload un fichier sur Supabase et retourne le lien de téléchargement"""
                        time_stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                        file_path = f"{time_stamp}_{filename}"

                        # 🔹 Upload vers Supabase Storage
                        resp = supabase.storage.from_("devis").upload(
                            file_path,
                            file,  # Fichier en bytes
                            file_options={
                                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
                            # Ajout du MIME type ici
                        )
                        # 🔹 Vérification des erreurs
                        if isinstance(resp, dict) and "error" in resp and resp["error"]:
                            return None, resp["error"]

                        # 🔹 Générer l'URL publique
                        url = supabase.storage.from_("devis").get_public_url(file_path)
                        return url, None

                    file_bytes = generate_word_doc()
                    file_url, error = upload_to_supabase(file_bytes, "mon_devis.docx")

                    if error:
                        self.download_button.disabled = True
                    else:
                        self.download_button.url = file_url
                        self.download_button.visible = True
                        self.download_button.update()

            # Si la TVA n'est pas active
            else:
                # Si l'IR est active
                if ir is True:
                    # generer le document word
                    def generate_word_doc():
                        doc = Document()
                        # Ajouter une image dans l'en-tête
                        section = doc.sections[0]
                        header = section.header
                        header_paragraph = header.paragraphs[0]
                        footer = section.footer
                        footer_paragraph = footer.paragraphs[0]

                        # Ajouter l'image dans l'en-tête
                        header_paragraph.add_run().add_picture("assets/images/header.jpg", width=Inches(6.5))
                        # Ajouter l'image dans le pied de page
                        footer_paragraph.add_run().add_picture("assets/images/footer.png", width=Inches(6.5))

                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        return buffer.getvalue()

                    def upload_to_supabase(file, filename):
                        """Upload un fichier sur Supabase et retourne le lien de téléchargement"""
                        time_stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                        file_path = f"{time_stamp}_{filename}"

                        # 🔹 Upload vers Supabase Storage
                        resp = supabase.storage.from_("devis").upload(
                            file_path,
                            file,  # Fichier en bytes
                            file_options={
                                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
                            # Ajout du MIME type ici
                        )
                        # 🔹 Vérification des erreurs
                        if isinstance(resp, dict) and "error" in resp and resp["error"]:
                            return None, resp["error"]

                        # 🔹 Générer l'URL publique
                        url = supabase.storage.from_("devis").get_public_url(file_path)
                        return url, None

                    file_bytes = generate_word_doc()
                    file_url, error = upload_to_supabase(file_bytes, "mon_devis.docx")

                    if error:
                        self.download_button.disabled = True
                    else:
                        self.download_button.url = file_url
                        self.download_button.visible = True
                        self.download_button.update()

                # Si l'IR n'est pos active
                else:
                    # generer le document word
                    def generate_word_doc():
                        doc = Document()
                        # Ajouter une image dans l'en-tête
                        section = doc.sections[0]
                        header = section.header
                        header_paragraph = header.paragraphs[0]
                        footer = section.footer
                        footer_paragraph = footer.paragraphs[0]

                        # Ajouter l'image dans l'en-tête
                        header_paragraph.add_run().add_picture("assets/images/header.jpg", width=Inches(6.5))
                        # Ajouter l'image dans le pied de page
                        footer_paragraph.add_run().add_picture("assets/images/footer.png", width=Inches(6.5))

                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        return buffer.getvalue()

                    def upload_to_supabase(file, filename):
                        """Upload un fichier sur Supabase et retourne le lien de téléchargement"""
                        time_stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                        file_path = f"{time_stamp}_{filename}"

                        # 🔹 Upload vers Supabase Storage
                        resp = supabase.storage.from_("devis").upload(
                            file_path,
                            file,  # Fichier en bytes
                            file_options={
                                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
                            # Ajout du MIME type ici
                        )
                        # 🔹 Vérification des erreurs
                        if isinstance(resp, dict) and "error" in resp and resp["error"]:
                            return None, resp["error"]

                        # 🔹 Générer l'URL publique
                        url = supabase.storage.from_("devis").get_public_url(file_path)
                        return url, None

                    file_bytes = generate_word_doc()
                    file_url, error = upload_to_supabase(file_bytes, "mon_devis.docx")

                    if error:
                        self.download_button.disabled = True
                    else:
                        self.download_button.url = file_url
                        self.download_button.visible = True
                        self.download_button.update()


