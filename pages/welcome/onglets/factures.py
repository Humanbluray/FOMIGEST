from utils import *
import flet as ft
import backend as be
import datetime
from utils.useful_functions import ajout_separateur, ecrire_en_lettres
from utils.constantes import *
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import io
from dotenv import load_dotenv
from supabase import create_client

load_dotenv()
url = os.environ.get("SUPABASE_URL")
key = os.environ.get("SUPABASE_KEY")
supabase = create_client(url, key)


class OneArticle(ft.Container):
    def __init__(self, cp: object, ref: str, des: str, qte: int, prix: int):
        super().__init__(
            padding=ft.padding.only(10, 3, 10, 3),
        )
        self.cp = cp
        self.ref = ref
        self.des = des
        self.prix = prix
        self.qte = qte

        self.art_ref = ft.TextField(**readonly_field_style, value=ref, width=140, label="Référence",)
        self.art_des = ft.TextField(**readonly_field_style, value=des, width=300, label="Nom pièce", tooltip=f"{des}")
        self.art_qte = ft.TextField(**readonly_field_style, value=f"{qte}", width=80, label="Qté")
        self.art_prix = ft.TextField(**readonly_field_style, value=f"{ajout_separateur(prix)}", width=100, label="Prix")

        self.content = ft.Row(
            controls=[self.art_ref, self.art_des, self.art_qte, self.art_prix],
            alignment=ft.MainAxisAlignment.SPACE_BETWEEN
        )


class Factures(ft.Container):
    def __init__(self, cp: object):
        super().__init__(expand=True)
        self.cp = cp
        self.search = ft.TextField(**search_field_style, width=300, prefix_icon="search", on_change=self.filter_datas)
        self.results = ft.Text("", size=12, font_family="Poppins Medium")
        self.table = ft.DataTable(
            **datatable_style,
            columns=[
                ft.DataColumn(ft.Text("")),
                ft.DataColumn(ft.Text("numero".upper())),
                ft.DataColumn(ft.Text("devis".upper())),
                ft.DataColumn(ft.Text("Montant".upper())),
                ft.DataColumn(ft.Text("Mt. payé".upper())),
                ft.DataColumn(ft.Text("reste".upper())),
                ft.DataColumn(ft.Text("")),
            ]
        )
        self.main_window = ft.Container(
            expand=True,
            padding=ft.padding.only(20, 15, 20, 15), border_radius=10, bgcolor="white",
            content=ft.Column(
                controls=[
                    ft.Container(
                        bgcolor="#f0f0f6", padding=10, border_radius=16,
                        content=ft.Row(
                            controls=[
                                ft.Icon(ft.icons.CREDIT_CARD, color="black"),
                                ft.Text("Factures".upper(), size=24, font_family="Poppins Bold"),
                            ]
                        )
                    ),
                    ft.Divider(height=1, color=ft.colors.TRANSPARENT),
                    ft.Row(
                        controls=[
                            ft.Row(
                                controls=[
                                    self.search, self.results
                                ]
                            ),
                        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
                    ),
                    ft.Divider(height=1, color=ft.colors.TRANSPARENT),
                    ft.ListView(expand=True, controls=[self.table])
                ]
            )
        )
        self.pay_facture = ft.Text("", size=12, font_family="Poppins Medium", color="white")
        self.pay_montant = ft.TextField(**readonly_date_style, width=170, label="Total facture", prefix_icon=ft.icons.ATTACH_MONEY_OUTLINED)
        self.pay_deja_paye = ft.TextField(**readonly_date_style, width=170, label="Déja payé", prefix_icon=ft.icons.MONETIZATION_ON_OUTLINED)
        self.pay_solde = ft.Checkbox(
            check_color=SECOND_COLOR, fill_color="#f0f0f6", label="Solde la facture",
            label_style=ft.TextStyle(size=12, font_family="Poppins Medium"),
            on_change=self.changement_solde
        )
        self.pay_mode = ft.Dropdown(
            **drop_style, width=200, label="Mode de paiement", prefix_icon=ft.icons.CREDIT_CARD_OUTLINED,
            options=[
                ft.dropdown.Option("virement".upper()),
                ft.dropdown.Option("espèces".upper()),
                ft.dropdown.Option("Orange Money".upper()),
                ft.dropdown.Option("MTN Mobile Money".upper()),
            ]
        )
        self.cp.dp_paiement.on_change = self.changement_date
        self.bt_select_date = ft.IconButton(
            ft.icons.CALENDAR_MONTH_OUTLINED, icon_color="black", scale=0.65,
            on_click=lambda _: self.cp.dp_paiement.pick_date()
        )
        self.pay_selected_date = ft.TextField(**readonly_date_style, label="Date", width=200, prefix_icon=ft.icons.EDIT_CALENDAR_OUTLINED)
        self.pay_button = AnyButton(
            FIRST_COLOR, "check", "Valider paiement", "white", 185,
            self.create_paiement
        )
        self.pay_a_regler = ft.TextField(**numbers_field_style, width=180, label="Montant règlement", prefix_icon=ft.icons.ATTACH_MONEY_OUTLINED)
        self.new_paiement_window = ft.Container(
            bgcolor="#f0f0f6", width=300, height=530,border_radius=16,
            padding=10, expand=True,
            shadow=ft.BoxShadow(spread_radius=5, blur_radius=15, color=ft.colors.BLACK38),
            scale=ft.transform.Scale(0),
            animate_scale=ft.Animation(300, ft.AnimationCurve.DECELERATE),
            content=ft.Column(
                expand=True,
                controls=[
                    ft.Container(
                        bgcolor="white", padding=ft.padding.only(10, 5, 10, 5), border_radius=16,
                        content=ft.Row(
                            controls=[
                                ft.Row(
                                    controls=[
                                        ft.Icon(ft.icons.NOTE_ADD_OUTLINED, "black"),
                                        ft.Text("Nouveau paiement".upper(), size=14, font_family="Poppins Bold")
                                    ]
                                ),
                                ft.IconButton("close", FIRST_COLOR, scale=0.6, bgcolor="#f2f2f2",
                                              on_click=self.close_new_paiement)
                            ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
                        )
                    ),
                    ft.Container(
                        bgcolor="white", padding=20, border_radius=16, expand=True,
                        content=ft.Column(
                            controls=[
                                ft.Container(
                                    padding=ft.padding.only(10, 3, 10, 3), border_radius=10, bgcolor=SECOND_COLOR,
                                    content=ft.Row([self.pay_facture], alignment=ft.MainAxisAlignment.CENTER)
                                ),
                                self.pay_montant, self.pay_deja_paye, self.pay_solde,
                                self.pay_a_regler, self.pay_mode,
                                ft.Stack(
                                    controls=[self.pay_selected_date, self.bt_select_date],
                                    alignment=ft.alignment.center_right
                                ),
                                self.pay_button
                            ], spacing=15
                        )
                    )
                ]
            )
        )

        # Edit window ...
        self.edit_bt_facture = ft.Container(padding=ft.padding.only(10, 3, 10, 3), border_radius=10)
        self.edit_nb_ref = ft.Text("", size=12, font_family="Poppins Medium", italic=True, color="grey")
        self.edit_table = ft.ListView(expand=True, divider_thickness=1, spacing=10)
        self.edit_num = ft.Text("", size=12, font_family="Poppins Medium", color="white")
        self.edit_dev = ft.Text("", size=12, font_family="Poppins Medium", color="white")
        self.edit_statut = ft.Text("", size=12, font_family="Poppins Medium", color="white")
        self.edit_icon_statut = ft.Icon(color="white", size=16)
        self.edit_montant = ft.TextField(**readonly_date_style, width=170, prefix_icon=ft.icons.PRICE_CHANGE_OUTLINED, label="Montant")
        self.edit_lettres = ft.TextField(**readonly_field_style, width=700, prefix_icon=ft.icons.LABEL_OUTLINED, label="Montant en lettres")
        self.edit_objet = ft.TextField(**readonly_date_style, width=700, label="Objet", prefix_icon=ft.icons.MAIL_OUTLINED)
        self.edit_bc = ft.TextField(**readonly_date_style, width=200, label="BC client", prefix_icon=ft.icons.NOTES_OUTLINED)
        self.edit_ov = ft.TextField(**readonly_date_style, width=200, label="OV", prefix_icon=ft.icons.NOTES)
        self.edit_delai = ft.TextField(**readonly_date_style, width=150, label="Délai livraison", prefix_icon=ft.icons.TIMELAPSE_OUTLINED)
        self.edit_remise = ft.TextField(**readonly_date_style, width=100, label="remise", value="0", prefix_icon=ft.icons.KEYBOARD_ARROW_DOWN_OUTLINED)
        self.ct_statut = ft.Container(
            padding=ft.padding.only(10, 3, 10, 3), border_radius=10,
            content=ft.Row(
                controls=[
                    self.edit_icon_statut,
                    self.edit_statut
                ],
                alignment=ft.MainAxisAlignment.CENTER
            )
        )
        self.bt_print_options = AnyButton(FIRST_COLOR, "print_outlined", "Imprimer facture", "white", 230, self.open_impression_window)
        self.edit_window = ft.Card(
            elevation=20, surface_tint_color="#f0f0f6", width=900, height=650,
            clip_behavior=ft.ClipBehavior.ANTI_ALIAS, shadow_color="black",
            scale=ft.transform.Scale(0), expand=True,
            animate_scale=ft.Animation(300, ft.AnimationCurve.DECELERATE),
            content=ft.Container(
                padding=10, bgcolor="#f0f0f6", expand=True,
                content=ft.Column(
                    controls=[
                        ft.Container(
                            bgcolor="white", padding=ft.padding.only(10, 5, 10, 5), border_radius=16,
                            content=ft.Row(
                                controls=[
                                    ft.Row(
                                        controls=[
                                            ft.Icon(ft.icons.EDIT_OUTLINED, "black"),
                                            ft.Text("Détails facture".upper(), size=14,
                                                    font_family="Poppins Bold")
                                        ]
                                    ),
                                    ft.IconButton("close", FIRST_COLOR, scale=0.6, bgcolor="#f2f2f2",
                                                  on_click=self.close_edit_window)
                                ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
                            )
                        ),
                        ft.Container(
                            bgcolor="white", padding=20, border_radius=16, expand=True,
                            content=ft.Column(
                                controls=[
                                    ft.Row(
                                        controls=[
                                            ft.Row(
                                                controls=[
                                                    ft.Row(
                                                        controls=[
                                                            ft.Text("Facture N°".upper(), size=12, font_family="Poppins Medium"),
                                                            ft.Container(
                                                                padding=ft.padding.only(10, 3, 10, 3), border_radius=10,
                                                                bgcolor=SECOND_COLOR, content=ft.Row(
                                                                    controls=[self.edit_num, ],
                                                                    alignment=ft.MainAxisAlignment.CENTER
                                                                )
                                                            )
                                                        ]
                                                    ),
                                                    ft.Row(
                                                        controls=[
                                                            ft.Text("Devis N°".upper(), size=12,
                                                                    font_family="Poppins Medium"),
                                                            ft.Container(
                                                                padding=ft.padding.only(10, 3, 10, 3), border_radius=10,
                                                                bgcolor=ft.colors.DEEP_ORANGE, content=ft.Row(
                                                                    controls=[self.edit_dev],
                                                                    alignment=ft.MainAxisAlignment.CENTER
                                                                )
                                                            )
                                                        ]
                                                    ),
                                                    ft.Row(
                                                        controls=[
                                                            ft.Text("Statut".upper(), size=12,font_family="Poppins Medium"),
                                                            self.ct_statut
                                                        ]
                                                    ),
                                                    self.edit_bt_facture
                                                ]
                                            ),
                                            self.edit_nb_ref
                                        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
                                    ),
                                    ft.Container(
                                        height=240, padding=ft.padding.only(10, 3, 10, 3),
                                        border_radius=16,
                                        expand=True,
                                        border=ft.border.all(1, "grey"), content=self.edit_table
                                    ),
                                    self.edit_objet,
                                    ft.Row(
                                        controls=[
                                            self.edit_ov, self.edit_delai, self.edit_remise, self.edit_bc,
                                        ]
                                    ),

                                    ft.Row([self.edit_montant, self.edit_lettres]),
                                    ft.Row([self.bt_print_options], spacing=20)
                                ]
                            )
                        )
                    ]
                )
            )
        )

        # Paiements window
        self.vp_numero = ft.Text(size=12, font_family="Poppins Medium", color="white")
        self.vp_statut = ft.Text(size=12, font_family="Poppins Medium", color="white")
        self.vp_table = ft.DataTable(
            **datatable_style,
            columns=[
                ft.DataColumn(ft.Text("date".upper())),
                ft.DataColumn(ft.Text("type".upper())),
                ft.DataColumn(ft.Text("montant".upper())),
            ]
        )
        self.vp_icone_statut = ft.Icon(size=16, color="white")
        self.vp_ct_statut =  ft.Container(
            border_radius=10,
            padding=ft.padding.only(10, 3, 10, 3),
            content=ft.Row(
                [
                    self.vp_icone_statut,
                    self.vp_statut
                ],
                alignment=ft.MainAxisAlignment.CENTER
            )
        )

        self.paiements_window = ft.Container(
            bgcolor="#f0f0f6", width=600, height=400,border_radius=16,
            padding=10, expand=True,
            shadow=ft.BoxShadow(spread_radius=5, blur_radius=15, color=ft.colors.BLACK38),
            scale=ft.transform.Scale(0),
            animate_scale=ft.Animation(300, ft.AnimationCurve.DECELERATE),
            content=ft.Column(
                expand=True,
                controls=[
                    ft.Container(
                        bgcolor="white", padding=ft.padding.only(10, 5, 10, 5), border_radius=16,
                        content=ft.Row(
                            controls=[
                                ft.Row(
                                    controls=[
                                        ft.Icon(ft.icons.NOTE_ADD_OUTLINED, "black"),
                                        ft.Text("Paiements".upper(), size=14,
                                                font_family="Poppins Medium")
                                    ]
                                ),
                                ft.IconButton("close", FIRST_COLOR, scale=0.6, bgcolor="#f2f2f2",
                                              on_click=self.close_paiements_window)
                            ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
                        )
                    ),
                    ft.Container(
                        bgcolor="white", padding=20, border_radius=16, expand=True,
                        content=ft.Column(
                            controls=[
                                ft.Row(
                                    controls=[
                                        ft.Row(
                                            controls=[
                                                ft.Text("Numéro de facture:".upper(), size=12,
                                                        font_family="Poppins Medium"),
                                                ft.Container(
                                                    bgcolor=SECOND_COLOR, border_radius=10,
                                                    padding=ft.padding.only(10, 3, 10, 3),
                                                    content=ft.Row([self.vp_numero],
                                                                   alignment=ft.MainAxisAlignment.CENTER)
                                                )
                                            ]
                                        ),
                                        ft.Row(
                                            controls=[
                                                ft.Text("Statut:".upper(), size=12,
                                                        font_family="Poppins Medium"),
                                                self.vp_ct_statut
                                            ]
                                        ),
                                    ]
                                ),
                                ft.Container(
                                    padding=ft.padding.only(10, 5, 10, 5), border_radius=10,
                                    border=ft.border.all(1, "grey"), expand=True,
                                    content=ft.ListView(expand=True, controls=[self.vp_table])
                                )
                            ]
                        )
                    )
                ]
            )
        )

        # options d'impression window
        self.regime = ft.RadioGroup(
            content=ft.Row(
                controls=[
                    ft.Radio(
                        **radio_style, value="S", label="Simplifié".upper()
                    ),
                    ft.Radio(
                        **radio_style, value="R", label="Réel".upper()
                    )
                ]
            )
        )
        self.tva = ft.Checkbox(
            label_style=ft.TextStyle(size=12, font_family="Poppins Medium"), active_color="white",
            check_color=FIRST_COLOR,
            label="TVA", label_position=ft.LabelPosition.RIGHT
        )
        self.ir = ft.Checkbox(
            label_style=ft.TextStyle(size=12, font_family="Poppins Medium"), active_color="white",
            check_color=FIRST_COLOR,
            label="IR", label_position=ft.LabelPosition.RIGHT
        )
        self.banque = ft.RadioGroup(
            content=ft.Row(
                controls=[
                    ft.Radio(
                        **radio_style, value="CCA", label="CCA".upper()
                    ),
                    ft.Radio(
                        **radio_style, value="AFRILAND", label="AFRILAND".upper()
                    )
                ]
            )
        )
        self.mask_button = AnyButton(
            "white", None, "Télécharger fichier", FIRST_COLOR, 280, None
        )
        self.download_button = AnyButton(
            FIRST_COLOR, "edit", "Télécharger fichier", "white", 280, None
        )
        self.download_button.visible = False
        self.mask_button.disabled = True

        self.impression_window = ft.Card(
            elevation=20, surface_tint_color="#f0f0f6", width=350, height=450,
            clip_behavior=ft.ClipBehavior.ANTI_ALIAS, shadow_color="black",
            scale=ft.transform.Scale(0), expand=True,
            animate_scale=ft.Animation(300, ft.AnimationCurve.DECELERATE),
            content=ft.Container(
                padding=10, bgcolor="#f0f0f6",
                content=ft.Container(
                    padding=10, border_radius=16, bgcolor="white",
                    content=ft.Column(
                        controls=[
                            ft.Column(
                                controls=[
                                    ft.Row(
                                        controls=[
                                            ft.Text("Options d'impressions".upper(), size=14,
                                                    font_family="Poppins Bold"),
                                            ft.IconButton("close", FIRST_COLOR, scale=0.6, bgcolor="#f2f2f2",
                                                          on_click=self.close_impression_window)
                                        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN
                                    ),
                                    ft.Divider(height=1, thickness=1),
                                ], spacing=0
                            ),
                            ft.Divider(height=3, color=ft.colors.TRANSPARENT),
                            ft.Column(
                                controls=[
                                    ft.Text("choix du régime".upper(), size=11, font_family="Poppins Bold"),
                                    ft.Divider(height=1, thickness=1),
                                ], spacing=0
                            ),
                            self.regime,
                            ft.Divider(height=3, color=ft.colors.TRANSPARENT),
                            ft.Column(
                                controls=[
                                    ft.Text("TVA et IR".upper(), size=11, font_family="Poppins Bold"),
                                    ft.Divider(height=1, thickness=1),
                                ], spacing=0
                            ),
                            ft.Row([self.tva, self.ir]),
                            ft.Column(
                                controls=[
                                    ft.Text("Choix de la banque".upper(), size=11, font_family="Poppins Bold"),
                                    ft.Divider(height=1, thickness=1),
                                ], spacing=0
                            ),
                            self.banque,
                            ft.Divider(height=1, color=ft.colors.TRANSPARENT),
                            AnyButton(
                                FIRST_COLOR, ft.icons.LOCAL_PRINTSHOP_OUTLINED, "Imprimer", "white",
                                280, self.imprimer_facture
                            ),
                            self.mask_button, self.download_button

                        ], horizontal_alignment=ft.CrossAxisAlignment.CENTER
                    )
                )
            )
        )

        self.content = ft.Stack(
            controls=[
                self.main_window, self.new_paiement_window, self.edit_window, self.impression_window, self.paiements_window
            ], alignment=ft.alignment.center
        )
        self.load_datas()

    def load_datas(self):
        datas = be.all_factures()
        self.results.value = f"{len(datas)} résultat(s)"

        for row in self.table.rows[:]:
            self.table.rows.remove(row)

        for data in datas:
            if data["mt_remise"] > data["regle"]:
                statut = "en cours"
                icone = ft.icons.INDETERMINATE_CHECK_BOX
                couleur = "red"
                pay_bt = CtButton(ft.icons.ADD_CARD_OUTLINED, ft.colors.BLACK45, "Paiement", data, self.open_new_paiement)
            else:
                statut = "soldée"
                icone = ft.icons.CHECK
                couleur = ft.colors.TEAL_500
                pay_bt = CtButton(None, ft.colors.BLACK45, "", data, None)

            self.table.rows.append(
                ft.DataRow(
                    cells=[
                        ft.DataCell(ft.Icon(icone, couleur, 20)),
                        ft.DataCell(ft.Text(data["numero"])),
                        ft.DataCell(ft.Text(data["devis"])),
                        ft.DataCell(ft.Text(ajout_separateur(data["mt_remise"]))),
                        ft.DataCell(ft.Text(ajout_separateur(data["regle"]))),
                        ft.DataCell(ft.Text(f"{ajout_separateur(data['mt_remise'] - data['regle'])}")),
                        ft.DataCell(
                            ft.Row(
                                controls=[
                                    CtButton(ft.icons.EDIT_OUTLINED, ft.colors.BLUE_300, "Voir details", data, self.open_edit_window),
                                    pay_bt,
                                    CtButton(ft.icons.LIST, ft.colors.BLACK45, "paiements", data, self.voir_paiements)
                                ], spacing=0, alignment=ft.MainAxisAlignment.END
                            )
                        ),
                    ]
                )
            )

    def filter_datas(self, e):
        datas = be.all_factures()
        search = self.search.value if self.search.value is not None else ""
        filtered_datas = list(filter(lambda x: search in x["numero"] or search in x["devis"] or search in x["nom_client"], datas))

        self.results.value = f"{len(filtered_datas)} résultat(s)"
        self.results.update()

        for row in self.table.rows[:]:
            self.table.rows.remove(row)

        for data in filtered_datas:
            if data["mt_remise"] > data["regle"]:
                statut = "en cours"
                icone = ft.icons.INDETERMINATE_CHECK_BOX
                couleur = "red"
                pay_bt = CtButton(ft.icons.ADD_CARD_OUTLINED, ft.colors.BLACK45, "Paiement", data, self.open_new_paiement)
            else:
                statut = "soldée"
                icone = ft.icons.CHECK
                couleur = ft.colors.TEAL_500
                pay_bt = CtButton(None, ft.colors.BLACK45, "", data, None)

            self.table.rows.append(
                ft.DataRow(
                    cells=[
                        ft.DataCell(ft.Icon(icone, couleur, 20)),
                        ft.DataCell(ft.Text(data["numero"])),
                        ft.DataCell(ft.Text(data["devis"])),
                        ft.DataCell(ft.Text(ajout_separateur(data["mt_remise"]))),
                        ft.DataCell(ft.Text(ajout_separateur(data["regle"]))),
                        ft.DataCell(ft.Text(f"{ajout_separateur(data['mt_remise'] - data['regle'])}")),
                        ft.DataCell(
                            ft.Row(
                                controls=[
                                    CtButton(ft.icons.EDIT_OUTLINED, ft.colors.BLUE_300, "Voir details", data, self.open_edit_window),
                                    pay_bt,
                                    CtButton(ft.icons.LIST, ft.colors.BLACK45, "paiements", data, self.voir_paiements)
                                ], spacing=0, alignment=ft.MainAxisAlignment.END
                            )
                        ),
                    ]
                )
            )

        self.table.update()

    def open_new_paiement(self, e):
        self.pay_facture.value = e.control.data["numero"]
        self.pay_facture.update()
        self.pay_deja_paye.value = f"{e.control.data['regle']}"
        self.pay_deja_paye.update()
        self.pay_montant.value = f"{e.control.data['mt_remise']}"
        self.pay_montant.update()
        self.new_paiement_window.scale = 1
        self.new_paiement_window.update()
        self.main_window.disabled = True
        self.main_window.opacity = 0.3
        self.main_window.update()

    def close_new_paiement(self, e):
        self.new_paiement_window.scale = 0
        self.new_paiement_window.update()
        self.main_window.disabled = False
        self.main_window.opacity = 1
        self.main_window.update()

    def create_paiement(self, e):
        if self.pay_a_regler.value is None or self.pay_a_regler.value == "":
            self.cp.box.title.value = "Erreur"
            self.cp.box.content.value = "Attention au montant"
            self.cp.box.open = True
            self.cp.box.update()

        elif int(self.pay_a_regler.value) > (int(self.pay_montant.value) - int(self.pay_deja_paye.value)):
            self.cp.box.title.value = "Erreur"
            self.cp.box.content.value = "Solde payé > au montant de la facture !"
            self.cp.box.open = True
            self.cp.box.update()

        else:
            be.add_reglement(
                self.pay_facture.value, int(self.pay_a_regler.value), self.pay_mode.value, self.pay_selected_date.value
            )
            self.cp.box.title.value = "Confirmé"
            self.cp.box.content.value = "Paiement sauvegardé"
            self.cp.box.open = True
            self.cp.box.update()

            for widget in (
                self.pay_deja_paye, self.pay_montant, self.pay_facture, self.pay_mode,
                self.pay_a_regler, self.pay_solde, self.pay_selected_date
            ):
                widget.value = None
                widget.update()

            self.new_paiement_window.scale = 0
            self.new_paiement_window.update()

            self.load_datas()
            self.table.update()
            self.results.update()

    def changement_date(self, e):
        self.pay_selected_date.value = str(self.cp.dp_paiement.value)[0:10]
        self.pay_selected_date.update()

    def changement_solde(self, e):
        if self.pay_solde.value:
            self.pay_a_regler.value = f"{int(self.pay_montant.value) - int(self.pay_deja_paye.value)}"
            self.pay_a_regler.update()
        else:
            self.pay_a_regler.value = None
            self.pay_a_regler.update()

    def open_edit_window(self, e):
        # remplissage des détails de la facture
        self.edit_num.value = f"{e.control.data['numero']}"
        self.edit_dev.value = f"{e.control.data['devis']}"

        if e.control.data["reste"] == 0:
            self.edit_icon_statut.name = ft.icons.CHECK
            self.edit_statut.value = "soldée".upper()
            self.ct_statut.bgcolor = "green"
        else:
            self.edit_icon_statut.name = ft.icons.INDETERMINATE_CHECK_BOX
            self.edit_statut.value = "En cours".upper()
            self.ct_statut.bgcolor = "red"

        self.edit_objet.value = e.control.data["objet"]
        self.edit_delai.value = e.control.data["delai"]
        self.edit_remise.value = f"{e.control.data['remise']}"
        self.edit_montant.value = f"{e.control.data['mt_remise']}"
        self.edit_lettres.value = ecrire_en_lettres(e.control.data["mt_remise"])
        self.edit_bc.value = e.control.data["bc_client"]
        self.edit_ov.value = e.control.data["ov"]

        self.edit_lettres.update()
        self.edit_bc.update()
        self.edit_ov.update()
        self.edit_remise.update()
        self.edit_montant.update()
        self.edit_delai.update()
        self.edit_objet.update()
        self.edit_statut.update()
        self.edit_icon_statut.update()
        self.edit_num.update()
        self.edit_dev.update()
        self.ct_statut.update()

        # Remplissage des détails de la table
        details = be.factures_details(e.control.data["numero"])
        for widget in self.edit_table.controls[:]:
            self.edit_table.controls.remove(widget)

        for detail in details:
            self.edit_table.controls.append(
                OneArticle(self, detail["reference"], detail["designation"], detail["qte"], detail["prix"])
            )

        self.edit_table.update()
        self.edit_window.scale = 1
        self.edit_window.update()
        self.main_window.disabled = True
        self.main_window.opacity = 0.3
        self.main_window.update()

    def close_edit_window(self, e):
        self.edit_window.scale = 0
        self.edit_window.update()
        self.main_window.disabled = False
        self.main_window.opacity = 1
        self.main_window.update()

    def open_impression_window(self, e):
        self.impression_window.scale = 1
        self.impression_window.update()
        self.edit_window.disabled = True
        self.edit_window.update()

    def close_impression_window(self, e):
        self.impression_window.scale = 0
        self.impression_window.update()
        self.edit_window.disabled = False
        self.edit_window.opacity = 1
        self.edit_window.update()

        self.regime.value = None
        self.tva.value = None
        self.ir.value = None
        self.banque.value = None
        self.mask_button.visible = True
        self.download_button.visible = False
        self.download_button.url = None

        for widget in (
                self.regime, self.tva, self.ir, self.banque, self.mask_button, self.download_button
        ):
            widget.update()

    def voir_paiements(self, e):
        self.vp_numero.value = e.control.data["numero"]

        if e.control.data["reste"] == 0:
            self.vp_statut.value = "soldée".upper()
            self.vp_icone_statut.name = ft.icons.CHECK
            self.vp_ct_statut.bgcolor = "green"
        else:
            self.vp_statut.value = "en cours".upper()
            self.vp_icone_statut.name = ft.icons.INDETERMINATE_CHECK_BOX
            self.vp_ct_statut.bgcolor = "red"

        for row in self.vp_table.rows[:]:
            self.vp_table.rows.remove(row)

        datas = be.all_reglements_by_facture(e.control.data["numero"])

        for data in datas:
            self.vp_table.rows.append(
                ft.DataRow(
                    cells=[
                        ft.DataCell(ft.Text(data["date"])),
                        ft.DataCell(ft.Text(data["type"])),
                        ft.DataCell(ft.Text(ajout_separateur(data["montant"]))),
                    ]
                )
            )

        self.vp_table.update()

        self.vp_numero.update()
        self.vp_statut.update()
        self.vp_icone_statut.update()
        self.vp_ct_statut.update()
        self.paiements_window.scale = 1
        self.paiements_window.update()
        self.main_window.disabled = True
        self.main_window.opacity = 0.3
        self.main_window.update()

    def close_paiements_window(self, e):
        self.paiements_window.scale = 0
        self.paiements_window.update()
        self.main_window.disabled = False
        self.main_window.opacity = 1
        self.main_window.update()

    def imprimer_facture(self, e):
        regime = self.regime.value
        tva = self.tva.value
        ir = self.ir.value
        banque = self.banque.value

        # Si la TVA est Active
        if tva and regime == "S":
            # Erreur
            self.cp.box.title.value = "Erreur"
            self.cp.box.content.value = "Pas de TVA dans le régime simplifié"
            self.cp.box.open = True
            self.cp.box.update()

        # Pour tout autre cas
        else:
            # generer le document word
            total_prix = 0
            def generate_word_doc():
                doc = Document()

                # Ajouter une image dans l'en-tête
                def header_and_footer():
                    section = doc.sections[0]

                    section.left_margin = Cm(1.5)  # Marge gauche
                    section.right_margin = Cm(1.5)  # Marge droite
                    section.top_margin = Cm(0.5)  # Marge haute
                    section.bottom_margin = Cm(0.5)  # Marge basse

                    header = section.header
                    header_paragraph = header.paragraphs[0]
                    footer = section.footer
                    footer_paragraph = footer.paragraphs[0]

                    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # Ajouter l'image dans l'en-tête
                    header_paragraph.add_run().add_picture("assets/images/header.png", width=Cm(18))
                    # Ajouter l'image dans le pied de page
                    footer_paragraph.add_run().add_picture("assets/images/footer.png", width=Cm(18))

                header_and_footer()

                # Créer un tableau principal avec une seule ligne et deux colonnes pour entete client
                main_table = doc.add_table(rows=1, cols=2)

                # Entête infos facture
                def entete_facture():
                    # Ajouter un tableau dans la première cellule
                    cell1 = main_table.cell(0, 0)
                    num_facture = self.edit_num.value
                    bc_client = be.show_info_factures(num_facture)["bc_client"]
                    ov = be.show_info_factures(num_facture)['ov']

                    table1 = cell1.add_table(rows=4, cols=1)  # Tableau avec 4 lignes et 1 colonne
                    cell1_1 = table1.cell(0, 0)
                    paragraph1 = cell1_1.paragraphs[0]
                    run1 = paragraph1.add_run("FACTURE")
                    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run1.font.name = "Arial Black"
                    run1.font.size = Pt(14)
                    run1.font.color.rgb = RGBColor(200, 0, 0)
                    cell1_1.paragraphs[0].paragraph_format.space_before = 10  # Pas d'espace avant le paragraphe
                    cell1_1.paragraphs[0].paragraph_format.space_after = 10  # Pas d'espace après le paragraphe
                    cell1_1.paragraphs[0].paragraph_format.line_spacing = Pt(20)  # Espacement entre les lignes réduit

                    cell1_2 = table1.cell(1, 0)
                    paragraph2 = cell1_2.paragraphs[0]
                    run2 = paragraph2.add_run(f"{num_facture}")
                    paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run2.font.name = "calibri"
                    run2.font.size = Pt(11)
                    cell1_2.paragraphs[0].paragraph_format.space_before = 10  # Pas d'espace avant le paragraphe
                    cell1_2.paragraphs[0].paragraph_format.space_after = 10  # Pas d'espace après le paragraphe
                    cell1_2.paragraphs[0].paragraph_format.line_spacing = Pt(15)  # Espacement entre les lignes réduit

                    cell1_3 = table1.cell(2, 0)
                    paragraph3 = cell1_3.paragraphs[0]
                    run3 = paragraph3.add_run(f"Suivant BC N° {bc_client}")
                    paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run3.font.name = "calibri"
                    run3.font.size = Pt(11)
                    cell1_3.paragraphs[0].paragraph_format.space_before = 10  # Pas d'espace avant le paragraphe
                    cell1_3.paragraphs[0].paragraph_format.space_after = 10  # Pas d'espace après le paragraphe
                    cell1_3.paragraphs[0].paragraph_format.line_spacing = Pt(15)  # Espacement entre les lignes réduit

                    if ov == "" or ov is None:
                        pass
                    else:
                        cell1_4 = table1.cell(3, 0)
                        paragraph4 = cell1_4.paragraphs[0]
                        run4 = paragraph4.add_run(f"OV° {ov}")
                        paragraph4.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run4.font.name = "calibri"
                        run4.font.size = Pt(11)
                        cell1_4.paragraphs[0].paragraph_format.space_before = 10  # Pas d'espace avant le paragraphe
                        cell1_4.paragraphs[0].paragraph_format.space_after = 10  # Pas d'espace après le paragraphe
                        cell1_4.paragraphs[0].paragraph_format.line_spacing = Pt(
                            15)  # Espacement entre les lignes réduit

                entete_facture()

                # Entete infos client
                def entete_infos_client():
                    # Ajouter un tableau dans la première cellule
                    cell1 = main_table.cell(0, 1)
                    client_id = be.show_info_devis(self.edit_dev.value)["client"]
                    client = be.infos_clients(client_id)["nom"]
                    contact = be.infos_clients(client_id)["contact"]
                    nui = be.infos_clients(client_id)["NUI"]
                    rc = be.infos_clients(client_id)["RC"]

                    table1 = cell1.add_table(rows=4, cols=1)  # Tableau avec 4 lignes et 1 colonne
                    cell1_1 = table1.cell(0, 0)
                    paragraph1 = cell1_1.paragraphs[0]
                    run1 = paragraph1.add_run(f"Client:   {client}")
                    paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run1.font.name = "calibri"
                    run1.font.size = Pt(11)
                    cell1_1.paragraphs[0].paragraph_format.space_before = 10  # Pas d'espace avant le paragraphe
                    cell1_1.paragraphs[0].paragraph_format.space_after = 10  # Pas d'espace après le paragraphe
                    cell1_1.paragraphs[0].paragraph_format.line_spacing = Pt(15)  # Espacement entre les lignes réduit

                    cell1_2 = table1.cell(1, 0)
                    paragraph2 = cell1_2.paragraphs[0]
                    run2 = paragraph2.add_run(f"Contact:   {contact}")
                    paragraph2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run2.font.name = "calibri"
                    run2.font.size = Pt(11)
                    cell1_2.paragraphs[0].paragraph_format.space_before = 10  # Pas d'espace avant le paragraphe
                    cell1_2.paragraphs[0].paragraph_format.space_after = 10  # Pas d'espace après le paragraphe
                    cell1_2.paragraphs[0].paragraph_format.line_spacing = Pt(15)

                    cell1_3 = table1.cell(2, 0)
                    paragraph3 = cell1_3.paragraphs[0]
                    run3 = paragraph3.add_run(f"NUI:   {nui}")
                    paragraph3.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run3.font.name = "calibri"
                    run3.font.size = Pt(11)
                    cell1_3.paragraphs[0].paragraph_format.space_before = 10  # Pas d'espace avant le paragraphe
                    cell1_3.paragraphs[0].paragraph_format.space_after = 10  # Pas d'espace après le paragraphe
                    cell1_3.paragraphs[0].paragraph_format.line_spacing = Pt(15)

                    cell1_4 = table1.cell(3, 0)
                    paragraph4 = cell1_4.paragraphs[0]
                    run4 = paragraph4.add_run(f"RC:   {rc}")
                    paragraph4.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run4.font.name = "calibri"
                    run4.font.size = Pt(12)
                    cell1_4.paragraphs[0].paragraph_format.space_before = 10  # Pas d'espace avant le paragraphe
                    cell1_4.paragraphs[0].paragraph_format.space_after = 10  # Pas d'espace après le paragraphe
                    cell1_4.paragraphs[0].paragraph_format.line_spacing = Pt(15)

                entete_infos_client()

                # Fonction d'ajout ligne
                def ajouter_ligne_grise(paragraph):
                    p = paragraph._element  # Récupérer l'élément XML du paragraphe
                    pPr = p.find("w:pPr", paragraph._element.nsmap)  # Chercher les propriétés du paragraphe

                    if pPr is None:
                        pPr = OxmlElement("w:pPr")
                        p.insert(0, pPr)

                    pbdr = OxmlElement("w:pBdr")  # Élément pour les bordures

                    # Définir la bordure inférieure (ligne grise)
                    bottom_border = OxmlElement("w:bottom")
                    bottom_border.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
                                      "single")  # Type de ligne (simple)
                    bottom_border.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz",
                                      "6")  # Épaisseur de la ligne (6 = 0.5 pt)
                    bottom_border.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space",
                                      "1")  # Espacement avec le texte
                    bottom_border.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color",
                                      "808080")  # Couleur grise en hexadécimal

                    pbdr.append(bottom_border)  # Ajouter la bordure aux propriétés du paragraphe
                    pPr.append(pbdr)  # Appliquer la bordure au paragraphe

                # Fonction d'écriture
                def draw_simple_paragraph(text: str, alignment, before: int, after: int, font_size: int, is_italic: bool, is_bold: bool):
                    details_pg = doc.add_paragraph()
                    details_pg.alignment = alignment
                    details_pg_format = details_pg.paragraph_format
                    details_pg_format.space_before = Pt(before)  # Espace avant le paragraphe
                    details_pg_format.space_after = Pt(after)
                    details_pg_format.line_spacing = 1
                    run_details_header = details_pg.add_run(text)  # Run 1 Separation
                    run_details_header.font.name = "calibri"
                    run_details_header.font.size = Pt(font_size)
                    run_details_header.italic = is_italic
                    run_details_header.bold = is_bold

                # Objet
                objet = be.show_info_factures(self.edit_num.value)['objet']

                if objet == "" or objet is None:
                    pass
                else:
                    draw_simple_paragraph(
                        f"Objet: {objet}", WD_PARAGRAPH_ALIGNMENT.LEFT,
                        10, 10, 11, False, False
                    )

                # References
                def draw_details_devis():
                    details = be.factures_details(self.edit_num.value)
                    longueur = len(details) + 1

                    # Créer un tableau avec 3 lignes et 6 colonnes
                    table = doc.add_table(rows=longueur, cols=6)

                    # Appliquer des bordures au tableau entier
                    # Ajouter des bordures noires à chaque cellule
                    def set_cell_border(cell):
                        """ Ajoute une bordure noire autour d'une cellule """
                        tc_pr = cell._element.get_or_add_tcPr()
                        borders = OxmlElement('w:tcBorders')

                        for border_name in ['top', 'left', 'bottom', 'right']:
                            border = OxmlElement(f'w:{border_name}')
                            border.set(qn('w:val'), 'single')  # Bordure simple
                            border.set(qn('w:sz'),
                                       '8')  # Taille de la bordure (plus épais pour meilleure visibilité)
                            border.set(qn('w:space'), '0')
                            border.set(qn('w:color'), '000000')  # Noir
                            borders.append(border)

                        tc_pr.append(borders)

                    # Appliquer les bordures à toutes les cellules
                    for row in table.rows:
                        for cell in row.cells:
                            set_cell_border(cell)

                    # Définir la largeur spécifique des colonnes (en centimètres)
                    column_widths = [Cm(1), Cm(12), Cm(1.5), Cm(1), Cm(2), Cm(2.5)]  # Largeurs des colonnes

                    # Appliquer les largeurs de colonnes
                    for col_idx, width in enumerate(column_widths):
                        for row in table.rows:
                            row.cells[col_idx].width = width  # Définir la largeur de la cellule

                    # Ajouter les en-têtes (ligne 0)
                    hdr_cells = table.rows[0].cells
                    headers = ["item", "Désignation", "Qté", "U", "Prix", "total (CFA)"]
                    for i, hdr in enumerate(headers):
                        paragraph = hdr_cells[i].paragraphs[0]
                        run = paragraph.add_run(hdr)
                        run.bold = True
                        run.font.name = "Calibri"
                        run.font.size = Pt(10)

                        # Centrer horizontalement
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        # Centrer verticalement

                    # Ajouter des données dans les 2 autres lignes (lignes 1 et 2) avec des nombres aléatoires
                    for i, row in enumerate(table.rows[1:]):  # i commence à 0 pour la première ligne de données
                        # Indice de la ligne (Ligne 1, Ligne 2, etc.)
                        details = be.factures_details(self.edit_num.value)
                        values = [f"{i}", str(details[i]["designation"]), str(details[i]["qte"]),
                                  str(details[i]["unite"]), str(ajout_separateur(details[i]["prix"])),
                                  str(ajout_separateur(details[i]["prix"]*details[i]["qte"]))
                        ]

                        for j, cell in enumerate(row.cells):
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run(values[j])
                            run.font.name = "Calibri"
                            run.font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrer le texte
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                draw_details_devis()

                # Montant total
                mt_total = 0
                details = be.factures_details(self.edit_num.value)
                for row in details:
                    mt_total += row["qte"]*row["prix"]

                def draw_montants():
                    # Ecrire le montant
                    draw_simple_paragraph(
                        f"Total:    {ajout_separateur(mt_total)}", WD_PARAGRAPH_ALIGNMENT.RIGHT, 10, 1,
                        11, False, False
                    )

                    # 1er cas TVA active et IR actif
                    if tva and ir:
                        # si la remise est nulle
                        if int(self.edit_remise.value) == 0:
                            mt_taxe = int(mt_total * TVA_VALUE)
                            mt_ttc = mt_total - mt_taxe
                            mt_ir = int(mt_total * IR_VALUE[regime])
                            mt_nap = mt_ttc - mt_ir

                            draw_simple_paragraph(
                                f"TVA:    {ajout_separateur(mt_taxe)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Montant TTC:    {ajout_separateur(mt_ttc)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"IR:    {ajout_separateur(mt_ir)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"NAP:    {ajout_separateur(mt_nap)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Facture Proforma arrêtée à la somme de:", WD_PARAGRAPH_ALIGNMENT.CENTER, 10, 1,
                                10, True, False
                            )
                            draw_simple_paragraph(
                                f"{ecrire_en_lettres(mt_nap)}".upper(), WD_PARAGRAPH_ALIGNMENT.CENTER, 1, 20,
                                11, False, True
                            )

                        # Si la remise est non nulle
                        else:
                            rem = int(mt_total * int(self.edit_remise.value)/100)
                            mt_remise = mt_total - rem
                            mt_taxe = int(mt_remise * TVA_VALUE)
                            mt_ttc = mt_remise - mt_taxe
                            mt_ir = int(mt_total * IR_VALUE[regime])
                            mt_nap = mt_ttc - mt_ir

                            draw_simple_paragraph(
                                f"Remise:    {self.edit_remise.value} %", WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Montant remisé:    {ajout_separateur(mt_remise)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"TVA:    {ajout_separateur(mt_taxe)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Montant TTC:    {ajout_separateur(mt_ttc)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"IR:    {ajout_separateur(mt_ir)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"NAP:    {ajout_separateur(mt_nap)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Facture Proforma arrêtée à la somme de:", WD_PARAGRAPH_ALIGNMENT.CENTER, 10, 1,
                                10, True, False
                            )
                            draw_simple_paragraph(
                                f"{ecrire_en_lettres(mt_nap)}".upper(), WD_PARAGRAPH_ALIGNMENT.CENTER, 1, 20,
                                11, False, True
                            )

                    # 2e cas: TVA inactive et IR actif
                    elif not tva and ir:

                        # si la remise est nulle
                        if int(self.edit_remise.value) == 0:
                            mt_ir = int(mt_total*IR_VALUE[regime])
                            mt_nap = int(mt_total - mt_ir)
                            draw_simple_paragraph(
                                f"IR:    {ajout_separateur(mt_ir)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"NAP:    {ajout_separateur(mt_nap)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Facture Proforma arrêtée à la somme de:", WD_PARAGRAPH_ALIGNMENT.CENTER, 10, 1,
                                10, True, False
                            )
                            draw_simple_paragraph(
                                f"{ecrire_en_lettres(mt_nap)}".upper(), WD_PARAGRAPH_ALIGNMENT.CENTER, 1, 20,
                                11, False, True
                            )

                        # Si la remise est non nulle
                        else:
                            mt_remise = be.show_info_devis(self.edit_num.value)['montant']
                            mt_ir = int(mt_remise * IR_VALUE[regime])
                            mt_nap = mt_remise - mt_ir
                            draw_simple_paragraph(
                                f"Remise:    {self.edit_remise.value} %", WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Montant Remisé:    {ajout_separateur(mt_remise)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"IR:    {ajout_separateur(mt_ir)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"NAP:    {ajout_separateur(mt_nap)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Facture Proforma arrêtée à la somme de:", WD_PARAGRAPH_ALIGNMENT.CENTER, 10, 1,
                                10, True, False
                            )
                            draw_simple_paragraph(
                                f"{ecrire_en_lettres(mt_nap)}".upper(), WD_PARAGRAPH_ALIGNMENT.CENTER, 1, 20,
                                11, False, True
                            )

                    # 3e Cas: TVA actif IR inactif
                    elif tva and not ir:
                        # si la remise est nulle
                        if int(self.edit_remise.value) == 0:
                            mt_taxe = int(mt_total * TVA_VALUE)
                            mt_ttc = mt_total - mt_taxe

                            draw_simple_paragraph(
                                f"TVA:    {ajout_separateur(mt_taxe)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Montant TTC:    {ajout_separateur(mt_ttc)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Facture Proforma arrêtée à la somme de:", WD_PARAGRAPH_ALIGNMENT.CENTER, 10, 1,
                                10, True, False
                            )
                            draw_simple_paragraph(
                                f"{ecrire_en_lettres(mt_ttc)}".upper(), WD_PARAGRAPH_ALIGNMENT.CENTER, 1, 20,
                                11, False, True
                            )

                        # Si la remise est non nulle
                        else:
                            rem = int(mt_total * int(self.edit_remise.value) / 100)
                            mt_remise = mt_total - rem
                            mt_taxe = int(mt_remise * TVA_VALUE)
                            mt_ttc = mt_remise - mt_taxe

                            draw_simple_paragraph(
                                f"Remise:    {self.edit_remise.value} %", WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Montant remisé:    {ajout_separateur(mt_remise)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"TVA:    {ajout_separateur(mt_taxe)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Montant TTC:    {ajout_separateur(mt_ttc)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Facture Proforma arrêtée à la somme de:", WD_PARAGRAPH_ALIGNMENT.CENTER, 10, 1,
                                10, True, False
                            )
                            draw_simple_paragraph(
                                f"{ecrire_en_lettres(mt_ttc)}".upper(), WD_PARAGRAPH_ALIGNMENT.CENTER, 1, 20,
                                11, False, True
                            )

                    # 4e cas TVA et IR inactifs
                    else:
                        # si la remise est nulle
                        if int(self.edit_remise.value) == 0:
                            draw_simple_paragraph(
                                f"Facture Proforma arrêtée à la somme de:", WD_PARAGRAPH_ALIGNMENT.CENTER, 10, 1,
                                10, True, False
                            )
                            draw_simple_paragraph(
                                f"{ecrire_en_lettres(mt_total)}".upper(), WD_PARAGRAPH_ALIGNMENT.CENTER, 1, 20,
                                11, False, True
                            )

                        # Si la remise est non nulle
                        else:
                            rem = int(mt_total * int(self.edit_remise.value) / 100)
                            mt_remise = mt_total - rem

                            draw_simple_paragraph(
                                f"Remise:    {self.edit_remise.value} %", WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Montant remisé:    {ajout_separateur(mt_remise)}",
                                WD_PARAGRAPH_ALIGNMENT.RIGHT, 1, 1,
                                11, False, False
                            )
                            draw_simple_paragraph(
                                f"Facture Proforma arrêtée à la somme de:", WD_PARAGRAPH_ALIGNMENT.CENTER, 10, 1,
                                10, True, False
                            )
                            draw_simple_paragraph(
                                f"{ecrire_en_lettres(mt_remise)}".upper(), WD_PARAGRAPH_ALIGNMENT.CENTER, 1, 20,
                                11, False, True
                            )

                draw_montants()

                # NB
                if be.show_info_devis(self.edit_dev.value)["note_bene"] == "" or \
                        be.show_info_devis(self.edit_dev.value)["note_bene"] is None:
                    pass

                else:
                    draw_simple_paragraph(
                        f"NB".upper(), WD_PARAGRAPH_ALIGNMENT.LEFT, 1, 3,
                        11, False, True
                    )
                    observations = be.show_info_devis(self.edit_num.value)["note_bene"].split(";")
                    if ";" in observations:
                        print(observations)
                        draw_simple_paragraph(
                            f"{observations}", WD_PARAGRAPH_ALIGNMENT.LEFT, 3, 3,
                            11, False, False
                        )
                    else:
                        divisions = be.show_info_devis(self.edit_num.value)["note_bene"].split(";")
                        for observ in divisions:
                            print(observ)
                            draw_simple_paragraph(
                                f"{observ}", WD_PARAGRAPH_ALIGNMENT.LEFT, 3, 3,
                                11, False, False
                            )

                pgf5 = doc.add_paragraph()
                ajouter_ligne_grise(pgf5)
                pgf5.paragraph_format.space_before = Pt(0)
                pgf5.paragraph_format.space_after = Pt(0)

                # Infos (delai lirvaison, point de livraison, paeiment, validité)
                def draw_other_infos():
                    pg1 = doc.add_paragraph()
                    # Délai de livraison
                    run1 = pg1.add_run("délai de livraison : ".upper())
                    run1.font.name = "calibri"
                    run1.font.size = Pt(10)
                    run1.font.color.rgb = RGBColor(175, 175, 175)
                    run2 = pg1.add_run(f"{be.show_info_devis(self.edit_dev.value)["delai"]}".upper())
                    run2.font.name = "calibri"
                    run2.font.size = Pt(10)
                    run2.font.color.rgb = RGBColor(0, 0, 0)

                    # point de livraison
                    run3 = pg1.add_run("                                 Point de livraison : ".upper())
                    run3.font.name = "calibri"
                    run3.font.size = Pt(10)
                    run3.font.color.rgb = RGBColor(175, 175, 175)
                    run4 = pg1.add_run(f"{be.show_info_devis(self.edit_dev.value)["point_liv"]}".upper())
                    run4.font.name = "calibri"
                    run4.font.size = Pt(10)
                    run4.font.color.rgb = RGBColor(0, 0, 0)
                    pg1.paragraph_format.space_before = Pt(1)
                    pg1.paragraph_format.space_after = Pt(3)

                    pg2 = doc.add_paragraph()
                    # Paiement
                    run5 = pg2.add_run("paiement : ".upper())
                    run5.font.name = "calibri"
                    run5.font.size = Pt(10)
                    run5.font.color.rgb = RGBColor(175, 175, 175)
                    run6 = pg2.add_run(f"{be.show_info_devis(self.edit_num.value)["paiement"]} Jours après dépôt de facture".upper())
                    run6.font.name = "calibri"
                    run6.font.size = Pt(10)
                    run6.font.color.rgb = RGBColor(0, 0, 0)
                    # Paiement
                    run7 = pg2.add_run("                    Validité de l'offre : ".upper())
                    run7.font.name = "calibri"
                    run7.font.size = Pt(10)
                    run7.font.color.rgb = RGBColor(175, 175, 175)
                    run8 = pg2.add_run(f"{be.show_info_devis(self.edit_num.value)["validite"]} mois".upper())
                    run8.font.name = "calibri"
                    run8.font.size = Pt(10)
                    run8.font.color.rgb = RGBColor(0, 0, 0)
                    pg1.paragraph_format.space_before = Pt(0)
                    pg2.paragraph_format.space_after = Pt(0)
                    ajouter_ligne_grise(pg2)

                # draw_other_infos()

                def draw_banque():
                    pg_banque = doc.add_paragraph()
                    run1 = pg_banque.add_run("Information bancaires : ".upper())
                    run1.font.name = "calibri"
                    run1.font.size = Pt(10)
                    run1.font.italic = True
                    run1.font.color.rgb = RGBColor(175, 175, 175)
                    pg_banque.paragraph_format.space_after = Pt(1)

                    draw_simple_paragraph(
                        f"Par virement à {ENTITE_BANQUE[banque]}, IBAN: {ENTITE_IBAN[banque]}".upper(),
                        WD_ALIGN_PARAGRAPH.LEFT, 3, 3, 11, False, False
                    )
                    draw_simple_paragraph(
                        f"Code swift: {ENTITE_SWIFT[banque]}, titulaire: FOMIDERC SARL".upper(),
                        WD_ALIGN_PARAGRAPH.LEFT, 3, 3, 11, False, False
                    )

                draw_banque()


                # Enregistrement du fichier
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer.getvalue()

            # Enregistre sur le bucket supabase
            def upload_to_supabase(file, filename):
                """Upload un fichier sur Supabase et retourne le lien de téléchargement"""
                time_stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                file_path = f"{time_stamp}_{filename}"

                # Upload vers Supabase Storage
                resp = supabase.storage.from_("devis").upload(
                    file_path,
                    file,  # Fichier en bytes
                    file_options={
                        "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
                    # Ajout du MIME type ici
                )
                # 🔹 Vérification des erreurs
                if isinstance(resp, dict) and "error" in resp and resp["error"]:
                    return None, resp["error"]

                # 🔹 Générer l'URL publique
                url = supabase.storage.from_("devis").get_public_url(file_path)
                return url, None

            file_bytes = generate_word_doc()
            file_url, error = upload_to_supabase(file_bytes, "mon_devis.docx")

            # En cas d'erreur
            if error:
                self.download_button.visible = False
                self.mask_button.visible = True

            # si pas d'erreur
            else:
                self.download_button.url = file_url
                self.download_button.visible = True
                self.mask_button.visible = False

            self.download_button.update()
            self.mask_button.update()
